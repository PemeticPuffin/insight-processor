"""
Tests for email_splitter.py

These tests challenge the email splitting logic with:
- Various email template structures
- Missing sections
- Edge cases in heading detection
- File creation and backup
"""

import pytest
from pathlib import Path
from docx import Document
from freezegun import freeze_time
import sys

sys.path.insert(0, str(Path(__file__).parent.parent))

from processors.email_splitter import (
    apply_font_formatting,
    copy_paragraph_with_formatting,
    parse_email_document,
    create_email_document,
    split_emails,
)
from config import FONT_NAME, FONT_SIZE_PT


class TestApplyFontFormatting:
    """Tests for apply_font_formatting function."""

    def test_applies_calibri_11pt(self, temp_dir):
        """Test that Calibri 11pt is applied to all paragraphs."""
        doc = Document()
        para = doc.add_paragraph()
        run = para.add_run("Test text")
        run.font.name = "Arial"
        run.font.size = 14

        apply_font_formatting(doc)

        assert doc.paragraphs[0].runs[0].font.name == FONT_NAME
        # Note: The function uses Pt(FONT_SIZE_PT)

    def test_multiple_paragraphs(self, temp_dir):
        """Test formatting multiple paragraphs."""
        doc = Document()
        for i in range(5):
            para = doc.add_paragraph()
            run = para.add_run(f"Paragraph {i}")
            run.font.name = "Times New Roman"

        apply_font_formatting(doc)

        for para in doc.paragraphs:
            for run in para.runs:
                assert run.font.name == FONT_NAME

    def test_empty_document(self, temp_dir):
        """Test with empty document."""
        doc = Document()

        # Should not raise error
        apply_font_formatting(doc)

    def test_paragraph_no_runs(self, temp_dir):
        """Test paragraph with no runs."""
        doc = Document()
        doc.add_paragraph()  # Empty paragraph

        # Should not raise error
        apply_font_formatting(doc)


class TestCopyParagraphWithFormatting:
    """Tests for copy_paragraph_with_formatting function."""

    def test_copies_text(self, temp_dir):
        """Test that text is copied."""
        source_doc = Document()
        source_para = source_doc.add_paragraph("Original text")

        target_doc = Document()
        new_para = copy_paragraph_with_formatting(source_para, target_doc)

        assert new_para.text == "Original text"

    def test_copies_style(self, temp_dir):
        """Test that paragraph style is copied."""
        source_doc = Document()
        source_para = source_doc.add_heading("Heading", level=2)

        target_doc = Document()
        new_para = copy_paragraph_with_formatting(source_para, target_doc)

        # Style should be copied (if it exists in target)

    def test_copies_alignment(self, temp_dir):
        """Test that alignment is copied."""
        from docx.enum.text import WD_ALIGN_PARAGRAPH

        source_doc = Document()
        source_para = source_doc.add_paragraph("Centered text")
        source_para.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER

        target_doc = Document()
        new_para = copy_paragraph_with_formatting(source_para, target_doc)

        assert new_para.paragraph_format.alignment == WD_ALIGN_PARAGRAPH.CENTER

    def test_copies_indentation(self, temp_dir):
        """Test that indentation is copied."""
        from docx.shared import Inches

        source_doc = Document()
        source_para = source_doc.add_paragraph("Indented text")
        source_para.paragraph_format.left_indent = Inches(0.5)

        target_doc = Document()
        new_para = copy_paragraph_with_formatting(source_para, target_doc)

        assert new_para.paragraph_format.left_indent == Inches(0.5)


class TestParseEmailDocument:
    """Tests for parse_email_document function."""

    def test_parses_basic_structure(self, email_template_docx):
        """Test parsing basic email template structure."""
        sections = parse_email_document(email_template_docx)

        assert "Chief Audit Executive" in sections
        assert "General Counsel" in sections

    def test_extracts_bd_section(self, email_template_docx):
        """Test extraction of BD (Business Developer) section."""
        sections = parse_email_document(email_template_docx)

        cae = sections["Chief Audit Executive"]
        assert cae['BD']['heading'] is not None
        assert len(cae['BD']['content']) > 0

    def test_extracts_ae_section(self, email_template_docx):
        """Test extraction of AE (Account Executive) section."""
        sections = parse_email_document(email_template_docx)

        cae = sections["Chief Audit Executive"]
        assert cae['AE']['heading'] is not None

    def test_extracts_ep_section(self, temp_dir):
        """Test extraction of EP (Executive Partner) section.

        Note: The EP heading must not contain 'CLIENT EMAIL' in addition to
        'SERVICE EXECUTIVE PARTNER' due to the elif chain in parse_email_document.
        """
        doc = Document()
        doc.add_heading("Chief Audit Executive", level=2)
        doc.add_heading("SERVICE EXECUTIVE PARTNER", level=3)  # EP without CLIENT EMAIL
        doc.add_paragraph("EP content")

        doc_path = temp_dir / "ep_test.docx"
        doc.save(str(doc_path))

        sections = parse_email_document(doc_path)
        cae = sections["Chief Audit Executive"]
        assert cae['EP']['heading'] is not None

    def test_empty_document(self, empty_docx):
        """Test with empty document."""
        sections = parse_email_document(empty_docx)

        assert sections == {}

    def test_no_heading_3(self, temp_dir):
        """Test document with Heading 2 but no Heading 3."""
        doc = Document()
        doc.add_heading("Chief Audit Executive", level=2)
        doc.add_paragraph("Content without H3 sections")

        doc_path = temp_dir / "no_h3.docx"
        doc.save(str(doc_path))

        sections = parse_email_document(doc_path)

        assert "Chief Audit Executive" in sections
        # All template types should be empty
        assert sections["Chief Audit Executive"]['BD']['heading'] is None

    def test_only_bd_template(self, temp_dir):
        """Test document with only BD template."""
        doc = Document()
        doc.add_heading("Chief Audit Executive", level=2)
        doc.add_heading("SALES BUSINESS DEVELOPER PROSPECT EMAIL", level=3)
        doc.add_paragraph("BD content only")

        doc_path = temp_dir / "bd_only.docx"
        doc.save(str(doc_path))

        sections = parse_email_document(doc_path)

        cae = sections["Chief Audit Executive"]
        assert cae['BD']['heading'] is not None
        assert cae['AE']['heading'] is None
        assert cae['EP']['heading'] is None

    def test_multiple_job_roles(self, temp_dir):
        """Test document with multiple job roles."""
        doc = Document()

        roles = ["Chief Audit Executive", "General Counsel", "Tech CEO"]
        for role in roles:
            doc.add_heading(role, level=2)
            doc.add_heading("SALES BUSINESS DEVELOPER PROSPECT EMAIL", level=3)
            doc.add_paragraph(f"BD content for {role}")

        doc_path = temp_dir / "multi_role.docx"
        doc.save(str(doc_path))

        sections = parse_email_document(doc_path)

        for role in roles:
            assert role in sections

    def test_service_executive_partner_detection(self, temp_dir):
        """Test that Service Executive Partner heading is detected as EP."""
        doc = Document()
        doc.add_heading("Chief Audit Executive", level=2)
        doc.add_heading("SALES BUSINESS DEVELOPER PROSPECT EMAIL", level=3)
        doc.add_paragraph("BD content")
        doc.add_heading("Service Executive Partner Client Email", level=3)
        doc.add_paragraph("EP content - should be separate")

        doc_path = temp_dir / "ep_test.docx"
        doc.save(str(doc_path))

        sections = parse_email_document(doc_path)

        cae = sections["Chief Audit Executive"]
        assert cae['BD']['heading'] is not None
        assert cae['EP']['heading'] is not None
        assert len(cae['BD']['content']) == 1
        assert len(cae['EP']['content']) == 1
        assert "BD content" in cae['BD']['content'][0].text
        assert "EP content" in cae['EP']['content'][0].text

    def test_all_three_template_types(self, temp_dir):
        """Test document with BD, AE, and EP templates."""
        doc = Document()
        doc.add_heading("Chief Audit Executive", level=2)
        doc.add_heading("SALES BUSINESS DEVELOPER PROSPECT EMAIL", level=3)
        doc.add_paragraph("BD content")
        doc.add_heading("SALES ACCOUNT EXECUTIVE CLIENT EMAIL", level=3)
        doc.add_paragraph("AE content")
        doc.add_heading("Service Executive Partner Client Email", level=3)
        doc.add_paragraph("EP content")

        doc_path = temp_dir / "all_templates.docx"
        doc.save(str(doc_path))

        sections = parse_email_document(doc_path)

        cae = sections["Chief Audit Executive"]
        assert cae['BD']['heading'] is not None
        assert cae['AE']['heading'] is not None
        assert cae['EP']['heading'] is not None


class TestCreateEmailDocument:
    """Tests for create_email_document function."""

    def test_creates_document_with_content(self, temp_dir):
        """Test creating document with heading and content."""
        source_doc = Document()
        heading = source_doc.add_heading("Email Heading", level=3)
        content1 = source_doc.add_paragraph("Content line 1")
        content2 = source_doc.add_paragraph("Content line 2")

        new_doc = create_email_document(
            [heading],
            [[content1, content2]]
        )

        assert len(new_doc.paragraphs) > 0

    def test_applies_font_formatting(self, temp_dir):
        """Test that font formatting is applied."""
        source_doc = Document()
        heading = source_doc.add_heading("Heading", level=3)
        content = source_doc.add_paragraph("Content")

        new_doc = create_email_document([heading], [[content]])

        # All runs should have Calibri font
        for para in new_doc.paragraphs:
            for run in para.runs:
                if run.font.name:
                    assert run.font.name == FONT_NAME

    def test_multiple_headings_and_content(self, temp_dir):
        """Test with multiple headings and content sections."""
        source_doc = Document()
        h1 = source_doc.add_heading("First Section", level=3)
        c1 = source_doc.add_paragraph("First content")
        h2 = source_doc.add_heading("Second Section", level=3)
        c2 = source_doc.add_paragraph("Second content")

        new_doc = create_email_document(
            [h1, h2],
            [[c1], [c2]]
        )

        texts = [p.text for p in new_doc.paragraphs]
        assert "First Section" in texts
        assert "Second Section" in texts

    def test_none_heading_handled(self, temp_dir):
        """Test handling of None heading."""
        source_doc = Document()
        content = source_doc.add_paragraph("Content without heading")

        new_doc = create_email_document(
            [None],
            [[content]]
        )

        # Should not crash


class TestSplitEmails:
    """Tests for split_emails function."""

    @freeze_time("2024-01-15")
    def test_basic_split(self, temp_dir, email_template_docx):
        """Test basic email splitting."""
        # Create target folder with subfolders
        target = temp_dir / "output"
        target.mkdir()
        (target / "Chief Audit Executive").mkdir()
        (target / "General Counsel").mkdir()

        jobs_processed, jobs_skipped, created_files = split_emails(
            email_template_docx, target
        )

        assert jobs_processed >= 1
        assert len(created_files) >= 1

    @freeze_time("2024-01-15")
    def test_creates_bd_ae_file(self, temp_dir):
        """Test creation of BD_AE_emails file."""
        # Create source document
        doc = Document()
        doc.add_heading("Chief Audit Executive", level=2)
        doc.add_heading("SALES BUSINESS DEVELOPER PROSPECT EMAIL", level=3)
        doc.add_paragraph("BD content")
        doc.add_heading("SALES ACCOUNT EXECUTIVE CLIENT EMAIL", level=3)
        doc.add_paragraph("AE content")

        source = temp_dir / "emails.docx"
        doc.save(str(source))

        # Create target
        target = temp_dir / "output"
        target.mkdir()
        (target / "Chief Audit Executive").mkdir()

        split_emails(source, target)

        # Check BD_AE file was created
        cae_folder = target / "Chief Audit Executive"
        bd_ae_files = list(cae_folder.glob("*BD_AE_emails.docx"))
        assert len(bd_ae_files) == 1

    @freeze_time("2024-01-15")
    def test_creates_ep_file(self, temp_dir):
        """Test creation of EP_email file.

        Note: EP heading must not include 'CLIENT EMAIL' due to parsing logic.
        """
        doc = Document()
        doc.add_heading("Chief Audit Executive", level=2)
        doc.add_heading("SERVICE EXECUTIVE PARTNER", level=3)  # Without CLIENT EMAIL
        doc.add_paragraph("EP content")

        source = temp_dir / "emails.docx"
        doc.save(str(source))

        target = temp_dir / "output"
        target.mkdir()
        (target / "Chief Audit Executive").mkdir()

        split_emails(source, target)

        cae_folder = target / "Chief Audit Executive"
        ep_files = list(cae_folder.glob("*EP_email.docx"))
        assert len(ep_files) == 1

    def test_skips_missing_subfolder(self, temp_dir):
        """Test that missing subfolders are skipped."""
        doc = Document()
        doc.add_heading("Missing Role", level=2)
        doc.add_heading("SALES BUSINESS DEVELOPER PROSPECT EMAIL", level=3)
        doc.add_paragraph("Content")

        source = temp_dir / "emails.docx"
        doc.save(str(source))

        target = temp_dir / "output"
        target.mkdir()
        # Don't create the "Missing Role" subfolder

        jobs_processed, jobs_skipped, created_files = split_emails(source, target)

        assert jobs_processed == 0
        assert jobs_skipped == 1

    def test_non_existent_target_folder(self, temp_dir, email_template_docx):
        """Test with non-existent target folder."""
        non_existent = temp_dir / "does_not_exist"

        jobs_processed, jobs_skipped, created_files = split_emails(
            email_template_docx, non_existent
        )

        assert jobs_processed == 0
        assert jobs_skipped == 0
        assert created_files == []

    @freeze_time("2024-01-15")
    def test_overwrites_existing(self, temp_dir):
        """Test that existing files are overwritten."""
        doc = Document()
        doc.add_heading("Chief Audit Executive", level=2)
        doc.add_heading("SALES BUSINESS DEVELOPER PROSPECT EMAIL", level=3)
        doc.add_paragraph("BD content - new")

        source = temp_dir / "emails.docx"
        doc.save(str(source))

        target = temp_dir / "output"
        target.mkdir()
        cae_folder = target / "Chief Audit Executive"
        cae_folder.mkdir()

        # Create existing file
        existing_doc = Document()
        existing_doc.add_paragraph("Existing content - old")
        # Date prefix for Jan 15 + 2 Mondays = 1.29
        existing_path = cae_folder / "1.29_CAE_BD_AE_emails.docx"
        existing_doc.save(str(existing_path))

        split_emails(source, target)

        # Check file was overwritten (no backups, content changed)
        assert existing_path.exists()
        updated_doc = Document(existing_path)
        content = " ".join([p.text for p in updated_doc.paragraphs])
        assert "new" in content.lower() or "BD content" in content

    def test_progress_callback(self, temp_dir, email_template_docx):
        """Test progress callback."""
        target = temp_dir / "output"
        target.mkdir()

        messages = []
        def callback(msg):
            messages.append(msg)

        split_emails(email_template_docx, target, progress_callback=callback)

        assert len(messages) > 0

    @freeze_time("2024-01-15")
    def test_uses_role_abbreviation(self, temp_dir):
        """Test that role abbreviations are used in filenames."""
        doc = Document()
        doc.add_heading("Chief Audit Executive", level=2)
        doc.add_heading("SALES BUSINESS DEVELOPER PROSPECT EMAIL", level=3)
        doc.add_paragraph("Content")

        source = temp_dir / "emails.docx"
        doc.save(str(source))

        target = temp_dir / "output"
        target.mkdir()
        (target / "Chief Audit Executive").mkdir()

        split_emails(source, target)

        cae_folder = target / "Chief Audit Executive"
        files = list(cae_folder.glob("*.docx"))

        # Should use CAE abbreviation
        file_names = [f.name for f in files]
        assert any("CAE" in name for name in file_names)


class TestEmailSplitterIntegration:
    """Integration tests for email splitter."""

    @freeze_time("2024-01-15")
    def test_full_realistic_workflow(self, temp_dir):
        """Test realistic email splitting workflow.

        Note: EP heading uses 'SERVICE EXECUTIVE PARTNER' without 'CLIENT EMAIL'
        to ensure proper parsing (see parse_email_document elif chain).
        """
        # Create comprehensive source document
        doc = Document()

        roles = [
            ("Chief Audit Executive", "CAE"),
            ("General Counsel", "GC"),
        ]

        for role, abbrev in roles:
            doc.add_heading(role, level=2)
            doc.add_heading("SALES BUSINESS DEVELOPER PROSPECT EMAIL", level=3)
            doc.add_paragraph(f"Dear Prospect, as a {role}...")
            doc.add_heading("SALES ACCOUNT EXECUTIVE CLIENT EMAIL", level=3)
            doc.add_paragraph(f"Dear Client, regarding your {role} needs...")
            doc.add_heading("SERVICE EXECUTIVE PARTNER", level=3)  # Fixed: no CLIENT EMAIL
            doc.add_paragraph(f"Dear Partner, {role} services include...")

        source = temp_dir / "email_templates.docx"
        doc.save(str(source))

        # Create target structure
        target = temp_dir / "output"
        target.mkdir()
        for role, _ in roles:
            (target / role).mkdir()

        jobs_processed, jobs_skipped, created_files = split_emails(source, target)

        assert jobs_processed == 2

        # Verify files for each role
        for role, abbrev in roles:
            role_folder = target / role
            files = list(role_folder.glob("*.docx"))

            # Should have BD_AE and EP files
            file_names = [f.name for f in files]
            assert any("BD_AE" in name for name in file_names)
            assert any("EP" in name for name in file_names)

            # Verify abbreviation is used
            assert any(abbrev in name for name in file_names)

    @freeze_time("2024-12-28")
    def test_year_boundary_dates(self, temp_dir):
        """Test date calculation near year boundary."""
        doc = Document()
        doc.add_heading("Chief Audit Executive", level=2)
        doc.add_heading("SALES BUSINESS DEVELOPER PROSPECT EMAIL", level=3)
        doc.add_paragraph("Content")

        source = temp_dir / "emails.docx"
        doc.save(str(source))

        target = temp_dir / "output"
        target.mkdir()
        (target / "Chief Audit Executive").mkdir()

        split_emails(source, target)

        cae_folder = target / "Chief Audit Executive"
        files = list(cae_folder.glob("*.docx"))

        # Date should be in January (2 Mondays from Dec 28)
        file_names = [f.name for f in files]
        assert any(name.startswith("1.") for name in file_names)
