"""
Tests for insight_splitter.py

These tests challenge the document splitting logic with:
- Empty documents
- Documents without headings
- Documents with many headings
- Hyperlink preservation
- Edge cases in heading detection
"""

import pytest
from pathlib import Path
from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import sys

sys.path.insert(0, str(Path(__file__).parent.parent))

from processors.insight_splitter import (
    is_heading_2,
    extract_sections,
    create_section_document,
    split_insights,
    get_heading_list,
)


class TestIsHeading2:
    """Tests for is_heading_2 function."""

    def test_heading_2_detected(self, temp_dir):
        """Test that Heading 2 style is correctly detected."""
        doc = Document()
        para = doc.add_heading("Test Heading", level=2)

        assert is_heading_2(para) == True

    def test_heading_1_not_detected(self, temp_dir):
        """Test that Heading 1 is not detected as Heading 2."""
        doc = Document()
        para = doc.add_heading("Test Heading", level=1)

        assert is_heading_2(para) == False

    def test_heading_3_not_detected(self, temp_dir):
        """Test that Heading 3 is not detected as Heading 2."""
        doc = Document()
        para = doc.add_heading("Test Heading", level=3)

        assert is_heading_2(para) == False

    def test_normal_paragraph_not_detected(self, temp_dir):
        """Test that normal paragraph is not Heading 2."""
        doc = Document()
        para = doc.add_paragraph("Normal text")

        assert is_heading_2(para) == False

    def test_title_not_detected(self, temp_dir):
        """Test that Title style is not Heading 2."""
        doc = Document()
        para = doc.add_paragraph("Title", style="Title")

        assert is_heading_2(para) == False

    def test_all_heading_levels(self, temp_dir):
        """Test all heading levels 0-9."""
        doc = Document()

        for level in range(10):
            try:
                para = doc.add_heading(f"Heading {level}", level=level)
                result = is_heading_2(para)

                if level == 2:
                    assert result == True
                else:
                    assert result == False
            except ValueError:
                # Some heading levels might not be valid
                pass


class TestExtractSections:
    """Tests for extract_sections function."""

    def test_extract_single_section(self, temp_dir):
        """Test extracting a single section."""
        doc = Document()
        doc.add_heading("Section One", level=2)
        doc.add_paragraph("Content paragraph 1")
        doc.add_paragraph("Content paragraph 2")

        sections = extract_sections(doc)

        assert len(sections) == 1
        assert sections[0]['heading'] == "Section One"
        assert len(sections[0]['content']) == 3  # Heading + 2 paragraphs

    def test_extract_multiple_sections(self, temp_dir):
        """Test extracting multiple sections."""
        doc = Document()

        doc.add_heading("Section One", level=2)
        doc.add_paragraph("Content 1")

        doc.add_heading("Section Two", level=2)
        doc.add_paragraph("Content 2")

        doc.add_heading("Section Three", level=2)
        doc.add_paragraph("Content 3")

        sections = extract_sections(doc)

        assert len(sections) == 3
        assert sections[0]['heading'] == "Section One"
        assert sections[1]['heading'] == "Section Two"
        assert sections[2]['heading'] == "Section Three"

    def test_extract_empty_document(self, temp_dir):
        """Test extraction from empty document."""
        doc = Document()

        sections = extract_sections(doc)

        assert sections == []

    def test_extract_no_heading_2(self, temp_dir):
        """Test document with no Heading 2."""
        doc = Document()
        doc.add_heading("Title", level=1)
        doc.add_paragraph("Content")
        doc.add_heading("Subheading", level=3)

        sections = extract_sections(doc)

        assert sections == []

    def test_extract_consecutive_headings(self, temp_dir):
        """Test consecutive Heading 2 elements without content."""
        doc = Document()
        doc.add_heading("Section One", level=2)
        doc.add_heading("Section Two", level=2)
        doc.add_heading("Section Three", level=2)

        sections = extract_sections(doc)

        assert len(sections) == 3
        # Each section has only its heading
        assert len(sections[0]['content']) == 1
        assert len(sections[1]['content']) == 1
        assert len(sections[2]['content']) == 1

    def test_content_before_first_heading_ignored(self, temp_dir):
        """Test that content before first Heading 2 is ignored."""
        doc = Document()
        doc.add_paragraph("This is before any heading")
        doc.add_paragraph("More preamble text")
        doc.add_heading("First Section", level=2)
        doc.add_paragraph("Section content")

        sections = extract_sections(doc)

        assert len(sections) == 1
        assert sections[0]['heading'] == "First Section"

    def test_mixed_heading_levels(self, temp_dir):
        """Test document with mixed heading levels."""
        doc = Document()
        doc.add_heading("Title", level=1)
        doc.add_heading("Section One", level=2)
        doc.add_paragraph("Content")
        doc.add_heading("Subsection", level=3)
        doc.add_paragraph("Sub content")
        doc.add_heading("Section Two", level=2)
        doc.add_paragraph("More content")

        sections = extract_sections(doc)

        assert len(sections) == 2
        # First section includes the H3 and its content
        assert sections[0]['heading'] == "Section One"
        assert sections[1]['heading'] == "Section Two"

    def test_extract_preserves_paragraph_objects(self, temp_dir):
        """Test that extracted content contains paragraph objects."""
        doc = Document()
        doc.add_heading("Section", level=2)
        doc.add_paragraph("Content")

        sections = extract_sections(doc)

        for para in sections[0]['content']:
            # Each item should be a paragraph object
            assert hasattr(para, 'text')
            assert hasattr(para, 'style')

    def test_heading_with_whitespace(self, temp_dir):
        """Test heading with leading/trailing whitespace."""
        doc = Document()
        para = doc.add_heading("  Section Name  ", level=2)
        doc.add_paragraph("Content")

        sections = extract_sections(doc)

        assert sections[0]['heading'] == "Section Name"  # Stripped

    def test_empty_heading_text(self, temp_dir):
        """Test Heading 2 with empty text."""
        doc = Document()
        doc.add_heading("", level=2)
        doc.add_paragraph("Content")
        doc.add_heading("Real Section", level=2)

        sections = extract_sections(doc)

        # Empty heading creates a section with empty string key
        assert len(sections) == 2
        assert sections[0]['heading'] == ""
        assert sections[1]['heading'] == "Real Section"


class TestCreateSectionDocument:
    """Tests for create_section_document function."""

    def test_creates_new_document(self, temp_dir):
        """Test that a new document is created."""
        doc = Document()
        doc.add_heading("Section", level=2)
        doc.add_paragraph("Content")

        sections = extract_sections(doc)
        new_doc = create_section_document(sections[0], doc)

        # Verify it's a Document by checking it has the expected attributes
        assert hasattr(new_doc, 'paragraphs')
        assert hasattr(new_doc, 'save')

    def test_preserves_content(self, temp_dir):
        """Test that content is preserved in new document."""
        doc = Document()
        doc.add_heading("My Section", level=2)
        doc.add_paragraph("First paragraph")
        doc.add_paragraph("Second paragraph")

        sections = extract_sections(doc)
        new_doc = create_section_document(sections[0], doc)

        texts = [p.text for p in new_doc.paragraphs]
        assert "My Section" in texts
        assert "First paragraph" in texts
        assert "Second paragraph" in texts

    def test_document_can_be_saved(self, temp_dir):
        """Test that created document can be saved."""
        doc = Document()
        doc.add_heading("Section", level=2)
        doc.add_paragraph("Content")

        sections = extract_sections(doc)
        new_doc = create_section_document(sections[0], doc)

        save_path = temp_dir / "output.docx"
        new_doc.save(str(save_path))

        assert save_path.exists()


class TestSplitInsights:
    """Tests for split_insights function."""

    def test_basic_split(self, temp_dir):
        """Test basic document splitting."""
        # Create source document
        doc = Document()
        doc.add_heading("Section One", level=2)
        doc.add_paragraph("Content for section one")
        doc.add_heading("Section Two", level=2)
        doc.add_paragraph("Content for section two")

        source_path = temp_dir / "source.docx"
        doc.save(str(source_path))

        output_dir = temp_dir / "output"
        output_dir.mkdir()

        files_created, sections_found, created_files = split_insights(
            source_path, output_dir
        )

        assert files_created == 2
        assert sections_found == 2
        assert len(created_files) == 2

    def test_creates_subfolders(self, temp_dir):
        """Test that subfolders are created for each section."""
        doc = Document()
        doc.add_heading("Chief Audit Executive", level=2)
        doc.add_paragraph("CAE content")
        doc.add_heading("General Counsel", level=2)
        doc.add_paragraph("GC content")

        source_path = temp_dir / "source.docx"
        doc.save(str(source_path))

        output_dir = temp_dir / "output"
        output_dir.mkdir()

        split_insights(source_path, output_dir, create_subfolders=True)

        # Known roles use canonical folder names
        assert (output_dir / "CAE").is_dir()
        assert (output_dir / "GC").is_dir()

    def test_no_heading_2_returns_zero(self, temp_dir):
        """Test document with no Heading 2."""
        doc = Document()
        doc.add_paragraph("Just normal text")
        doc.add_heading("Title", level=1)

        source_path = temp_dir / "source.docx"
        doc.save(str(source_path))

        output_dir = temp_dir / "output"
        output_dir.mkdir()

        files_created, sections_found, created_files = split_insights(
            source_path, output_dir
        )

        assert files_created == 0
        assert sections_found == 0
        assert created_files == []

    def test_empty_document(self, empty_docx, temp_dir):
        """Test with empty document."""
        output_dir = temp_dir / "output"
        output_dir.mkdir()

        files_created, sections_found, created_files = split_insights(
            empty_docx, output_dir
        )

        assert files_created == 0
        assert sections_found == 0

    def test_create_subfolders_false(self, temp_dir):
        """Test with create_subfolders=False skips missing folders."""
        doc = Document()
        doc.add_heading("Missing Folder", level=2)
        doc.add_paragraph("Content")

        source_path = temp_dir / "source.docx"
        doc.save(str(source_path))

        output_dir = temp_dir / "output"
        output_dir.mkdir()
        # Don't create the subfolder

        files_created, sections_found, created_files = split_insights(
            source_path, output_dir, create_subfolders=False
        )

        assert files_created == 0  # Skipped because folder doesn't exist
        assert sections_found == 1

    def test_progress_callback(self, temp_dir):
        """Test that progress callback is called."""
        doc = Document()
        doc.add_heading("Section", level=2)
        doc.add_paragraph("Content")

        source_path = temp_dir / "source.docx"
        doc.save(str(source_path))

        output_dir = temp_dir / "output"
        output_dir.mkdir()

        messages = []
        def callback(msg):
            messages.append(msg)

        split_insights(source_path, output_dir, progress_callback=callback)

        assert len(messages) > 0

    def test_many_sections(self, temp_dir):
        """Test document with many sections."""
        doc = Document()

        for i in range(20):
            doc.add_heading(f"Section {i}", level=2)
            doc.add_paragraph(f"Content for section {i}")

        source_path = temp_dir / "source.docx"
        doc.save(str(source_path))

        output_dir = temp_dir / "output"
        output_dir.mkdir()

        files_created, sections_found, created_files = split_insights(
            source_path, output_dir
        )

        assert files_created == 20
        assert sections_found == 20

    def test_output_filename_matches_heading(self, temp_dir):
        """Test that output files are named after headings."""
        doc = Document()
        doc.add_heading("My Section Name", level=2)
        doc.add_paragraph("Content")

        source_path = temp_dir / "source.docx"
        doc.save(str(source_path))

        output_dir = temp_dir / "output"
        output_dir.mkdir()

        split_insights(source_path, output_dir)

        # Unknown roles keep spaces in folder name, filename keeps original heading
        expected_file = output_dir / "My Section Name" / "My Section Name.docx"
        assert expected_file.exists()


class TestGetHeadingList:
    """Tests for get_heading_list function."""

    def test_returns_heading_texts(self, temp_dir):
        """Test that all Heading 2 texts are returned."""
        doc = Document()
        doc.add_heading("First Heading", level=2)
        doc.add_heading("Second Heading", level=2)
        doc.add_heading("Third Heading", level=2)

        source_path = temp_dir / "source.docx"
        doc.save(str(source_path))

        headings = get_heading_list(source_path)

        assert headings == ["First Heading", "Second Heading", "Third Heading"]

    def test_ignores_other_headings(self, temp_dir):
        """Test that non-Heading 2 styles are ignored."""
        doc = Document()
        doc.add_heading("Title", level=1)
        doc.add_heading("Section", level=2)
        doc.add_heading("Subsection", level=3)

        source_path = temp_dir / "source.docx"
        doc.save(str(source_path))

        headings = get_heading_list(source_path)

        assert headings == ["Section"]

    def test_empty_document(self, empty_docx):
        """Test with empty document."""
        headings = get_heading_list(empty_docx)

        assert headings == []

    def test_strips_whitespace(self, temp_dir):
        """Test that heading text is stripped."""
        doc = Document()
        doc.add_heading("  Padded Heading  ", level=2)

        source_path = temp_dir / "source.docx"
        doc.save(str(source_path))

        headings = get_heading_list(source_path)

        assert headings == ["Padded Heading"]


class TestInsightSplitterIntegration:
    """Integration tests for insight splitter."""

    def test_full_workflow_with_real_roles(self, temp_dir):
        """Test full splitting with actual role names."""
        doc = Document()

        roles = [
            "Chief Audit Executive",
            "General Counsel",
            "Tech CEO",
        ]

        for role in roles:
            doc.add_heading(role, level=2)
            doc.add_paragraph(f"Insight content for {role}")
            doc.add_paragraph(f"Additional details for {role}")

        source_path = temp_dir / "insights.docx"
        doc.save(str(source_path))

        output_dir = temp_dir / "split_output"
        output_dir.mkdir()

        files_created, sections_found, created_files = split_insights(
            source_path, output_dir
        )

        assert files_created == 3
        assert sections_found == 3

        # Verify each folder and file exists
        # Known roles use canonical folder names, filenames keep original heading
        canonical_folders = [
            "CAE",
            "GC",
            "TCEO",
        ]
        for role, canonical in zip(roles, canonical_folders):
            folder = output_dir / canonical
            assert folder.is_dir()

            doc_file = folder / f"{role}.docx"
            assert doc_file.exists()

            # Verify content
            created_doc = Document(doc_file)
            texts = [p.text for p in created_doc.paragraphs]
            assert role in texts[0]

    def test_get_headings_then_split(self, temp_dir):
        """Test getting heading list before splitting."""
        doc = Document()
        doc.add_heading("Section A", level=2)
        doc.add_heading("Section B", level=2)

        source_path = temp_dir / "source.docx"
        doc.save(str(source_path))

        # First, get the heading list
        headings = get_heading_list(source_path)
        assert len(headings) == 2

        # Then split
        output_dir = temp_dir / "output"
        output_dir.mkdir()

        files_created, sections_found, _ = split_insights(source_path, output_dir)

        assert files_created == len(headings)
