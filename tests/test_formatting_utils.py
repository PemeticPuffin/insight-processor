"""
Tests for formatting_utils.py

These tests challenge formatting functions with:
- Complex paragraph structures
- Edge cases in intro line detection
- Various list formats
- Font and style preservation
"""

import pytest
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import sys
from pathlib import Path

sys.path.insert(0, str(Path(__file__).parent.parent))

from utils.formatting_utils import (
    set_font_calibri_11pt,
    bold_paragraph,
    apply_bullet_style,
    add_spacing_after,
    replace_cxo_in_document,
)
from config import FONT_NAME, FONT_SIZE_PT


class TestSetFontCalibri11pt:
    """Tests for set_font_calibri_11pt function."""

    def test_basic_paragraph(self, temp_dir):
        """Test setting font on a basic paragraph."""
        doc = Document()
        para = doc.add_paragraph("Test paragraph")

        set_font_calibri_11pt(para)

        for run in para.runs:
            assert run.font.name == FONT_NAME
            assert run.font.size == Pt(FONT_SIZE_PT)

    def test_multiple_runs(self, temp_dir):
        """Test setting font on paragraph with multiple runs."""
        doc = Document()
        para = doc.add_paragraph()
        para.add_run("First part ")
        para.add_run("second part ")
        para.add_run("third part")

        set_font_calibri_11pt(para)

        assert len(para.runs) == 3
        for run in para.runs:
            assert run.font.name == FONT_NAME
            assert run.font.size == Pt(FONT_SIZE_PT)

    def test_empty_paragraph(self, temp_dir):
        """Test setting font on empty paragraph."""
        doc = Document()
        para = doc.add_paragraph("")

        # Should not raise error
        set_font_calibri_11pt(para)

    def test_paragraph_with_no_runs(self, temp_dir):
        """Test paragraph with no runs."""
        doc = Document()
        para = doc.add_paragraph()

        # Should not raise error
        set_font_calibri_11pt(para)
        assert len(para.runs) == 0

    def test_preserves_bold(self, temp_dir):
        """Test that bold formatting is preserved."""
        doc = Document()
        para = doc.add_paragraph()
        run = para.add_run("Bold text")
        run.bold = True

        set_font_calibri_11pt(para)

        # Bold should be preserved
        assert para.runs[0].bold == True

    def test_preserves_italic(self, temp_dir):
        """Test that italic formatting is preserved."""
        doc = Document()
        para = doc.add_paragraph()
        run = para.add_run("Italic text")
        run.italic = True

        set_font_calibri_11pt(para)

        assert para.runs[0].italic == True

    def test_overwrites_different_font(self, temp_dir):
        """Test that different font is overwritten."""
        doc = Document()
        para = doc.add_paragraph()
        run = para.add_run("Different font")
        run.font.name = "Arial"
        run.font.size = Pt(14)

        set_font_calibri_11pt(para)

        assert para.runs[0].font.name == FONT_NAME
        assert para.runs[0].font.size == Pt(FONT_SIZE_PT)


class TestBoldParagraph:
    """Tests for bold_paragraph function."""

    def test_basic_bold(self, temp_dir):
        """Test making a paragraph bold."""
        doc = Document()
        para = doc.add_paragraph("Text to bold")

        bold_paragraph(para)

        for run in para.runs:
            assert run.bold == True

    def test_multiple_runs_all_bold(self, temp_dir):
        """Test that all runs become bold."""
        doc = Document()
        para = doc.add_paragraph()
        para.add_run("First ")
        para.add_run("Second ")
        para.add_run("Third")

        bold_paragraph(para)

        for run in para.runs:
            assert run.bold == True

    def test_already_bold_stays_bold(self, temp_dir):
        """Test that already bold text stays bold."""
        doc = Document()
        para = doc.add_paragraph()
        run = para.add_run("Already bold")
        run.bold = True

        bold_paragraph(para)

        assert para.runs[0].bold == True

    def test_empty_paragraph(self, temp_dir):
        """Test bolding empty paragraph."""
        doc = Document()
        para = doc.add_paragraph("")

        # Should not raise error
        bold_paragraph(para)

    def test_preserves_other_formatting(self, temp_dir):
        """Test that other formatting is preserved."""
        doc = Document()
        para = doc.add_paragraph()
        run = para.add_run("Italic text")
        run.italic = True
        run.underline = True

        bold_paragraph(para)

        assert para.runs[0].bold == True
        assert para.runs[0].italic == True
        assert para.runs[0].underline == True


class TestApplyBulletStyle:
    """Tests for apply_bullet_style function."""

    def test_applies_bullet_formatting(self, temp_dir):
        """Test that bullet character and formatting is applied."""
        doc = Document()
        para = doc.add_paragraph("Bullet item")

        apply_bullet_style(para)

        # Should have Normal style (not List Bullet which can have numbering)
        assert para.style.name == "Normal"
        # Should start with bullet character
        assert para.text.startswith('\u2022')  # •

    def test_multiple_paragraphs(self, temp_dir):
        """Test applying bullets to multiple paragraphs."""
        doc = Document()
        paras = [
            doc.add_paragraph("Item 1"),
            doc.add_paragraph("Item 2"),
            doc.add_paragraph("Item 3"),
        ]

        for para in paras:
            apply_bullet_style(para)

        for para in paras:
            assert para.style.name == "Normal"
            assert para.text.startswith('\u2022')  # •

    def test_preserves_text_content(self, temp_dir):
        """Test that text content is preserved (with bullet prefix)."""
        doc = Document()
        original_text = "Important bullet point"
        para = doc.add_paragraph(original_text)

        apply_bullet_style(para)

        # Text should contain original content with bullet prefix
        assert original_text in para.text
        assert para.text.startswith('\u2022')  # •


class TestAddSpacingAfter:
    """Tests for add_spacing_after function."""

    def test_default_spacing(self, temp_dir):
        """Test default 12pt spacing."""
        doc = Document()
        para = doc.add_paragraph("Test paragraph")

        add_spacing_after(para)

        assert para.paragraph_format.space_after == Pt(12)

    def test_custom_spacing(self, temp_dir):
        """Test custom spacing values."""
        doc = Document()
        para = doc.add_paragraph("Test paragraph")

        add_spacing_after(para, points=24)

        assert para.paragraph_format.space_after == Pt(24)

    def test_zero_spacing(self, temp_dir):
        """Test zero spacing."""
        doc = Document()
        para = doc.add_paragraph("Test paragraph")

        add_spacing_after(para, points=0)

        assert para.paragraph_format.space_after == Pt(0)

    def test_large_spacing(self, temp_dir):
        """Test large spacing value."""
        doc = Document()
        para = doc.add_paragraph("Test paragraph")

        add_spacing_after(para, points=72)

        assert para.paragraph_format.space_after == Pt(72)


class TestFormattingUtilsIntegration:
    """Integration tests for formatting utilities."""

    def test_full_formatting_workflow(self, temp_dir):
        """Test typical formatting workflow."""
        doc = Document()
        para = doc.add_paragraph("1. Important point")

        set_font_calibri_11pt(para)
        apply_bullet_style(para)

        assert para.style.name == "Normal"
        assert para.text.startswith('\u2022')  # •

    def test_format_header_then_content(self, temp_dir):
        """Test formatting header and content differently."""
        doc = Document()

        header = doc.add_paragraph("Why This Matters")
        set_font_calibri_11pt(header)
        bold_paragraph(header)

        content = doc.add_paragraph("- Key insight point")
        set_font_calibri_11pt(content)
        apply_bullet_style(content)

        assert header.runs[0].bold == True
        assert content.style.name == "Normal"
        assert content.text.startswith('\u2022')  # •


class TestReplaceCxoInDocument:
    """Tests for replace_cxo_in_document function."""

    def test_replaces_cxo_in_single_run(self):
        """Test basic CXO replacement in a single run."""
        doc = Document()
        para = doc.add_paragraph()
        para.add_run("Hello CXO, welcome.")

        replace_cxo_in_document(doc, "Chief Information Officer")

        assert "CXO" not in para.text
        assert "Chief Information Officer" in para.text

    def test_replaces_cxo_in_multiple_paragraphs(self):
        """Test replacement across multiple paragraphs."""
        doc = Document()
        doc.add_paragraph().add_run("Dear CXO")
        doc.add_paragraph().add_run("As a CXO, you should know")
        doc.add_paragraph().add_run("No placeholder here")

        replace_cxo_in_document(doc, "Chief Financial Officer")

        for para in doc.paragraphs:
            assert "CXO" not in para.text

        assert "Chief Financial Officer" in doc.paragraphs[0].text
        assert "Chief Financial Officer" in doc.paragraphs[1].text
        assert "No placeholder here" == doc.paragraphs[2].text

    def test_replaces_multiple_cxo_in_one_run(self):
        """Test replacement of multiple CXO occurrences in one run."""
        doc = Document()
        para = doc.add_paragraph()
        para.add_run("CXO and CXO agenda")

        replace_cxo_in_document(doc, "Chief Marketing Officer")

        assert "CXO" not in para.text
        assert para.text.count("Chief Marketing Officer") == 2

    def test_no_cxo_no_change(self):
        """Test that documents without CXO are unaffected."""
        doc = Document()
        para = doc.add_paragraph()
        para.add_run("Regular content with no placeholder")

        replace_cxo_in_document(doc, "Chief Audit Executive")

        assert para.text == "Regular content with no placeholder"

    def test_preserves_run_formatting(self):
        """Test that bold/italic formatting on the run is preserved after replacement."""
        doc = Document()
        para = doc.add_paragraph()
        run = para.add_run("Attention CXO")
        run.bold = True

        replace_cxo_in_document(doc, "General Counsel")

        assert para.runs[0].bold == True
        assert "General Counsel" in para.text

    @pytest.mark.parametrize("variant", ["CXO", "cxo", "Cxo", "CxO", "cXo", "CXo"])
    def test_case_insensitive_replacement(self, variant):
        """Test that all case variants of CXO are replaced."""
        doc = Document()
        para = doc.add_paragraph()
        para.add_run(f"Hello {variant} welcome")

        replace_cxo_in_document(doc, "Chief Information Officer")

        assert variant not in para.text
        assert "Chief Information Officer" in para.text

    def test_empty_document(self):
        """Test with a document that has no paragraphs with text."""
        doc = Document()
        doc.add_paragraph()  # empty paragraph

        # Should not raise
        replace_cxo_in_document(doc, "Chief Information Officer")
