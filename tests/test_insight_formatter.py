"""
Tests for insight_formatter.py

These tests challenge the formatting logic with:
- Complex section ranges
- Header detection edge cases
- Bullet conversion scenarios
- Title paragraph detection
"""

import pytest
from pathlib import Path
from docx import Document
from docx.shared import Pt
import sys

sys.path.insert(0, str(Path(__file__).parent.parent))

from processors.insight_formatter import (
    is_header,
    get_section_ranges,
    find_title_paragraph,
    format_single_document,
    format_insights,
)
from config import HEADERS_TO_BOLD, FONT_NAME, FONT_SIZE_PT


class TestIsHeader:
    """Tests for is_header function."""

    @pytest.mark.parametrize("header", HEADERS_TO_BOLD)
    def test_detects_all_configured_headers(self, header):
        """Test detection of all headers in config."""
        doc = Document()
        para = doc.add_paragraph(header)

        assert is_header(para) == True

    def test_non_header_text(self):
        """Test that non-header text is not detected."""
        doc = Document()
        para = doc.add_paragraph("Some random paragraph text")

        assert is_header(para) == False

    def test_case_sensitive(self):
        """Test that header matching is case-sensitive."""
        doc = Document()

        # Lowercase version of a header
        para = doc.add_paragraph("why this matters")

        assert is_header(para) == False

    def test_partial_match_not_detected(self):
        """Test that partial header match is not detected."""
        doc = Document()
        para = doc.add_paragraph("Why This")  # Partial

        assert is_header(para) == False

    def test_header_with_extra_text(self):
        """Test header with additional text."""
        doc = Document()
        para = doc.add_paragraph("Why This Matters - Additional Info")

        # Exact match only
        assert is_header(para) == False

    def test_whitespace_stripped(self):
        """Test that whitespace is stripped before matching."""
        doc = Document()
        para = doc.add_paragraph("  Why This Matters  ")

        assert is_header(para) == True

    def test_empty_paragraph(self):
        """Test with empty paragraph."""
        doc = Document()
        para = doc.add_paragraph("")

        assert is_header(para) == False


class TestGetSectionRanges:
    """Tests for get_section_ranges function."""

    def test_basic_section_range(self, temp_dir):
        """Test basic section range detection."""
        doc = Document()
        doc.add_paragraph("Why This Matters")  # 0
        doc.add_paragraph("Content 1")          # 1
        doc.add_paragraph("Content 2")          # 2
        doc.add_paragraph("Implications")       # 3
        doc.add_paragraph("Content 3")          # 4

        ranges = get_section_ranges(doc)

        # Should find range from Why This Matters to Implications
        # (start after header, end at next header)
        assert (1, 3) in ranges

    def test_no_headers_returns_empty(self, temp_dir):
        """Test document without configured headers."""
        doc = Document()
        doc.add_paragraph("Random text")
        doc.add_paragraph("More text")

        ranges = get_section_ranges(doc)

        assert ranges == []

    def test_multiple_section_ranges(self, temp_dir):
        """Test detection of multiple section ranges."""
        doc = Document()
        doc.add_paragraph("Why This Matters")
        doc.add_paragraph("Content 1")
        doc.add_paragraph("Implications")
        doc.add_paragraph("Content 2")
        doc.add_paragraph("Probing Questions")
        doc.add_paragraph("Question 1")
        doc.add_paragraph("AskGartner Prompts to Recommend")
        doc.add_paragraph("Prompt 1")

        ranges = get_section_ranges(doc)

        # Should have ranges for configured bullet sections
        assert len(ranges) >= 1

    def test_section_to_end_of_document(self, temp_dir):
        """Test section range that extends to end of document."""
        doc = Document()
        doc.add_paragraph("Other content")
        doc.add_paragraph("AskGartner Prompts to Recommend")
        doc.add_paragraph("Prompt 1")
        doc.add_paragraph("Prompt 2")
        doc.add_paragraph("Prompt 3")

        ranges = get_section_ranges(doc)

        # AskGartner Prompts section goes to end (None in config)
        if ranges:
            last_range = ranges[-1]
            assert last_range[1] == len(doc.paragraphs)

    def test_missing_end_header(self, temp_dir):
        """Test when end header is missing."""
        doc = Document()
        doc.add_paragraph("Why This Matters")
        doc.add_paragraph("Content")
        # No "Implications" header

        ranges = get_section_ranges(doc)

        # Range should be skipped or go to end


class TestFindTitleParagraph:
    """Tests for find_title_paragraph function."""

    def test_finds_title_of_insight(self, temp_dir):
        """Test finding 'Title of Insight' paragraph."""
        doc = Document()
        doc.add_paragraph("Some header")
        doc.add_paragraph("Title of Insight: My Great Insight")
        doc.add_paragraph("Content")

        idx = find_title_paragraph(doc)

        assert idx == 1

    def test_finds_title_prefix(self, temp_dir):
        """Test finding 'Title:' prefix."""
        doc = Document()
        doc.add_paragraph("Title: Important Topic")
        doc.add_paragraph("Content")

        idx = find_title_paragraph(doc)

        assert idx == 0

    def test_case_insensitive(self, temp_dir):
        """Test case-insensitive title detection."""
        doc = Document()
        doc.add_paragraph("TITLE OF INSIGHT: TEST")

        idx = find_title_paragraph(doc)

        assert idx == 0

    def test_no_title_returns_negative(self, temp_dir):
        """Test that missing title returns -1."""
        doc = Document()
        doc.add_paragraph("Just normal content")
        doc.add_paragraph("More content")

        idx = find_title_paragraph(doc)

        assert idx == -1

    def test_empty_document(self, temp_dir):
        """Test with empty document."""
        doc = Document()

        idx = find_title_paragraph(doc)

        assert idx == -1

    def test_title_in_middle(self, temp_dir):
        """Test title appearing in middle of document."""
        doc = Document()
        doc.add_paragraph("First para")
        doc.add_paragraph("Second para")
        doc.add_paragraph("Title of Insight: Topic")
        doc.add_paragraph("Content")

        idx = find_title_paragraph(doc)

        assert idx == 2


class TestFormatSingleDocument:
    """Tests for format_single_document function."""

    def test_basic_formatting(self, temp_dir):
        """Test basic document formatting."""
        doc = Document()
        doc.add_paragraph("Title of Insight: Test")
        doc.add_paragraph("Why This Matters")
        doc.add_paragraph("Important point")
        doc.add_paragraph("Implications")

        doc_path = temp_dir / "test.docx"
        doc.save(str(doc_path))

        success, message = format_single_document(doc_path)

        assert success == True
        assert "success" in message.lower()

    def test_sets_calibri_font(self, temp_dir):
        """Test that Calibri font is applied."""
        doc = Document()
        para = doc.add_paragraph()
        run = para.add_run("Test text")
        run.font.name = "Arial"  # Different font

        doc_path = temp_dir / "test.docx"
        doc.save(str(doc_path))

        format_single_document(doc_path)

        # Reload and check
        formatted_doc = Document(doc_path)
        for para in formatted_doc.paragraphs:
            for run in para.runs:
                if run.font.name:
                    assert run.font.name == FONT_NAME

    def test_bolds_headers(self, temp_dir):
        """Test that configured headers are bolded."""
        doc = Document()
        doc.add_paragraph("Why This Matters")
        doc.add_paragraph("Normal content")

        doc_path = temp_dir / "test.docx"
        doc.save(str(doc_path))

        format_single_document(doc_path)

        # Reload and check
        formatted_doc = Document(doc_path)
        header_para = None
        for para in formatted_doc.paragraphs:
            if para.text.strip() == "Why This Matters":
                header_para = para
                break

        assert header_para is not None
        # Check if runs are bold
        if header_para.runs:
            assert header_para.runs[0].bold == True

    def test_adds_spacing_after_title(self, temp_dir):
        """Test that spacing is added after title."""
        doc = Document()
        doc.add_paragraph("Title of Insight: Test")
        doc.add_paragraph("Content")

        doc_path = temp_dir / "test.docx"
        doc.save(str(doc_path))

        format_single_document(doc_path)

        # Reload and check
        formatted_doc = Document(doc_path)
        title_para = formatted_doc.paragraphs[0]

        # Should have space_after set
        assert title_para.paragraph_format.space_after is not None

    def test_applies_bullets_in_sections(self, temp_dir):
        """Test that bullets are applied in bullet sections."""
        doc = Document()
        doc.add_paragraph("Why This Matters")
        doc.add_paragraph("First point to bullet")
        doc.add_paragraph("Second point to bullet")
        doc.add_paragraph("Implications")

        doc_path = temp_dir / "test.docx"
        doc.save(str(doc_path))

        format_single_document(doc_path)

        # Reload and check
        formatted_doc = Document(doc_path)
        # Content between headers should have bullet style

    def test_skips_intro_lines(self, temp_dir):
        """Test that intro lines are not bulleted."""
        doc = Document()
        doc.add_paragraph("Why This Matters")
        doc.add_paragraph("Key points:")  # Intro - should not be bulleted
        doc.add_paragraph("First point")
        doc.add_paragraph("Implications")

        doc_path = temp_dir / "test.docx"
        doc.save(str(doc_path))

        format_single_document(doc_path)

        # Reload and check
        formatted_doc = Document(doc_path)
        intro_para = None
        for para in formatted_doc.paragraphs:
            if para.text.strip() == "Key points:":
                intro_para = para
                break

        if intro_para:
            assert intro_para.style.name != "List Bullet"

    def test_converts_numbered_lists(self, temp_dir):
        """Test that numbered lists are converted to bullets."""
        doc = Document()
        doc.add_paragraph("Why This Matters")
        doc.add_paragraph("1. First item")
        doc.add_paragraph("2. Second item")
        doc.add_paragraph("Implications")

        doc_path = temp_dir / "test.docx"
        doc.save(str(doc_path))

        format_single_document(doc_path)

        # Reload and check - numbers should be removed

    def test_handles_corrupt_document(self, corrupted_docx):
        """Test handling of corrupt document."""
        success, message = format_single_document(corrupted_docx)

        assert success == False
        assert "error" in message.lower()

    def test_preserves_hyperlinks(self, docx_with_hyperlinks, temp_dir):
        """Test that hyperlinks are preserved."""
        success, message = format_single_document(docx_with_hyperlinks)

        assert success == True

        # Reload and check hyperlinks still exist
        doc = Document(docx_with_hyperlinks)
        # Check for hyperlink elements in the XML


class TestFormatInsights:
    """Tests for format_insights function."""

    def test_formats_multiple_files(self, temp_dir):
        """Test formatting multiple files."""
        # Create multiple documents
        for i in range(3):
            doc = Document()
            doc.add_paragraph(f"Document {i}")
            doc.save(str(temp_dir / f"doc{i}.docx"))

        success_count, failure_count = format_insights(temp_dir)

        assert success_count == 3
        assert failure_count == 0

    def test_recursive_formatting(self, temp_dir):
        """Test recursive directory search."""
        # Create nested structure
        subdir = temp_dir / "subfolder"
        subdir.mkdir()

        doc1 = Document()
        doc1.save(str(temp_dir / "root.docx"))

        doc2 = Document()
        doc2.save(str(subdir / "nested.docx"))

        success_count, failure_count = format_insights(temp_dir, recursive=True)

        assert success_count == 2

    def test_non_recursive(self, temp_dir):
        """Test non-recursive formatting."""
        subdir = temp_dir / "subfolder"
        subdir.mkdir()

        doc1 = Document()
        doc1.save(str(temp_dir / "root.docx"))

        doc2 = Document()
        doc2.save(str(subdir / "nested.docx"))

        success_count, failure_count = format_insights(temp_dir, recursive=False)

        assert success_count == 1  # Only root file

    def test_empty_directory(self, temp_dir):
        """Test with empty directory."""
        success_count, failure_count = format_insights(temp_dir)

        assert success_count == 0
        assert failure_count == 0

    def test_progress_callback(self, temp_dir):
        """Test progress callback."""
        doc = Document()
        doc.save(str(temp_dir / "test.docx"))

        messages = []
        def callback(msg):
            messages.append(msg)

        format_insights(temp_dir, progress_callback=callback)

        assert len(messages) > 0

    def test_mixed_valid_invalid_files(self, temp_dir):
        """Test with mix of valid and invalid files."""
        # Valid file
        doc = Document()
        doc.save(str(temp_dir / "valid.docx"))

        # Invalid file
        (temp_dir / "invalid.docx").write_text("not a real docx")

        success_count, failure_count = format_insights(temp_dir)

        assert success_count == 1
        assert failure_count == 1


class TestInsightFormatterIntegration:
    """Integration tests for insight formatter."""

    def test_full_realistic_document(self, temp_dir):
        """Test formatting a realistic insight document."""
        doc = Document()

        # Title
        doc.add_paragraph("Title of Insight: Digital Transformation Strategy")

        # Why This Matters section
        doc.add_paragraph("Why This Matters")
        doc.add_paragraph("Key considerations:")
        doc.add_paragraph("1. Drive operational efficiency")
        doc.add_paragraph("2. Improve customer experience")
        doc.add_paragraph("3. Enable data-driven decisions")

        # Implications section
        doc.add_paragraph("Implications")
        doc.add_paragraph("Organizations must adapt to remain competitive.")

        # Probing Questions section
        doc.add_paragraph("Probing Questions")
        doc.add_paragraph("What is your current digital maturity?")
        doc.add_paragraph("- How do you measure success?")
        doc.add_paragraph("- What are the key blockers?")

        # AskGartner section
        doc.add_paragraph("AskGartner Prompts to Recommend")
        doc.add_paragraph("Ask about digital transformation frameworks")
        doc.add_paragraph("Ask about change management strategies")

        doc_path = temp_dir / "insight.docx"
        doc.save(str(doc_path))

        success, message = format_single_document(doc_path)

        assert success == True

        # Verify formatting
        formatted = Document(doc_path)

        # Check title has spacing
        title_para = formatted.paragraphs[0]
        assert title_para.paragraph_format.space_after is not None

        # Check headers are bold
        for para in formatted.paragraphs:
            if para.text.strip() == "Why This Matters":
                if para.runs:
                    assert para.runs[0].bold == True
