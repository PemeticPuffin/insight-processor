"""
Tests for qc/validators.py

These tests challenge the validation logic with:
- Valid and invalid documents
- Missing sections
- Edge cases in heading detection
- Hyperlink validation
"""

import pytest
from pathlib import Path
from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import sys

sys.path.insert(0, str(Path(__file__).parent.parent))

from qc.validators import (
    validate_insights_document,
    validate_emails_document,
    validate_split_files,
    validate_formatting,
    validate_hyperlinks,
)


class TestValidateInsightsDocument:
    """Tests for validate_insights_document function."""

    def test_valid_document(self, temp_dir):
        """Test validation of a valid insights document."""
        doc = Document()
        doc.add_heading("First Section", level=2)
        doc.add_paragraph("Content")
        doc.add_heading("Second Section", level=2)
        doc.add_paragraph("More content")

        doc_path = temp_dir / "valid.docx"
        doc.save(str(doc_path))

        success, messages = validate_insights_document(doc_path)

        assert success == True
        assert any("Found" in msg and "Heading 2" in msg for msg in messages)

    def test_no_heading_2(self, temp_dir):
        """Test document without Heading 2."""
        doc = Document()
        doc.add_heading("Title", level=1)
        doc.add_paragraph("Content")
        doc.add_heading("Subheading", level=3)

        doc_path = temp_dir / "no_h2.docx"
        doc.save(str(doc_path))

        success, messages = validate_insights_document(doc_path)

        assert success == False
        assert any("No Heading 2" in msg for msg in messages)

    def test_file_not_found(self, temp_dir):
        """Test with non-existent file."""
        non_existent = temp_dir / "does_not_exist.docx"

        success, messages = validate_insights_document(non_existent)

        assert success == False
        assert any("not found" in msg.lower() for msg in messages)

    def test_corrupted_document(self, corrupted_docx):
        """Test with corrupted document."""
        success, messages = validate_insights_document(corrupted_docx)

        assert success == False
        assert any("failed" in msg.lower() for msg in messages)

    def test_empty_document(self, empty_docx):
        """Test with empty document."""
        success, messages = validate_insights_document(empty_docx)

        assert success == False
        assert any("No Heading 2" in msg for msg in messages)

    def test_many_headings_truncated(self, temp_dir):
        """Test that many headings are truncated in output."""
        doc = Document()
        for i in range(15):
            doc.add_heading(f"Section {i}", level=2)
            doc.add_paragraph("Content")

        doc_path = temp_dir / "many_headings.docx"
        doc.save(str(doc_path))

        success, messages = validate_insights_document(doc_path)

        assert success == True
        # Should show first 10 and "... and X more"
        assert any("more" in msg for msg in messages)

    def test_reports_heading_names(self, temp_dir):
        """Test that heading names are reported."""
        doc = Document()
        doc.add_heading("My Custom Section", level=2)
        doc.add_paragraph("Content")

        doc_path = temp_dir / "test.docx"
        doc.save(str(doc_path))

        success, messages = validate_insights_document(doc_path)

        assert any("My Custom Section" in msg for msg in messages)


class TestValidateEmailsDocument:
    """Tests for validate_emails_document function."""

    def test_valid_email_document(self, email_template_docx):
        """Test validation of valid email template."""
        success, messages = validate_emails_document(email_template_docx)

        assert success == True
        assert any("job role" in msg.lower() for msg in messages)

    def test_no_heading_2(self, temp_dir):
        """Test document without Heading 2 (job roles)."""
        doc = Document()
        doc.add_paragraph("Just content")
        doc.add_heading("Subheading", level=3)

        doc_path = temp_dir / "no_roles.docx"
        doc.save(str(doc_path))

        success, messages = validate_emails_document(doc_path)

        assert success == False
        assert any("No Heading 2" in msg for msg in messages)

    def test_counts_bd_sections(self, temp_dir):
        """Test that BD sections are counted."""
        doc = Document()
        doc.add_heading("Role 1", level=2)
        doc.add_heading("SALES BUSINESS DEVELOPER PROSPECT EMAIL", level=3)
        doc.add_paragraph("Content")
        doc.add_heading("Role 2", level=2)
        doc.add_heading("PROSPECT EMAIL", level=3)  # Also BD
        doc.add_paragraph("Content")

        doc_path = temp_dir / "bd_test.docx"
        doc.save(str(doc_path))

        success, messages = validate_emails_document(doc_path)

        assert any("BD sections: 2" in msg for msg in messages)

    def test_counts_ae_sections(self, temp_dir):
        """Test that AE sections are counted."""
        doc = Document()
        doc.add_heading("Role", level=2)
        doc.add_heading("SALES ACCOUNT EXECUTIVE CLIENT EMAIL", level=3)
        doc.add_paragraph("Content")

        doc_path = temp_dir / "ae_test.docx"
        doc.save(str(doc_path))

        success, messages = validate_emails_document(doc_path)

        assert any("AE sections: 1" in msg for msg in messages)

    def test_counts_ep_sections(self, temp_dir):
        """Test that EP sections are counted."""
        doc = Document()
        doc.add_heading("Role", level=2)
        doc.add_heading("SERVICE EXECUTIVE PARTNER CLIENT EMAIL", level=3)
        doc.add_paragraph("Content")

        doc_path = temp_dir / "ep_test.docx"
        doc.save(str(doc_path))

        success, messages = validate_emails_document(doc_path)

        assert any("EP sections: 1" in msg for msg in messages)

    def test_warns_missing_template_types(self, temp_dir):
        """Test warnings for missing template types."""
        doc = Document()
        doc.add_heading("Role", level=2)
        doc.add_heading("SALES BUSINESS DEVELOPER PROSPECT EMAIL", level=3)
        doc.add_paragraph("Only BD")

        doc_path = temp_dir / "bd_only.docx"
        doc.save(str(doc_path))

        success, messages = validate_emails_document(doc_path)

        assert any("Warning" in msg and "AE" in msg for msg in messages)
        assert any("Warning" in msg and "EP" in msg for msg in messages)

    def test_file_not_found(self, temp_dir):
        """Test with non-existent file."""
        non_existent = temp_dir / "does_not_exist.docx"

        success, messages = validate_emails_document(non_existent)

        assert success == False


class TestValidateSplitFiles:
    """Tests for validate_split_files function."""

    def test_all_files_present(self, temp_dir):
        """Test when all expected files are present."""
        headings = ["Section1", "Section2", "Section3"]

        for heading in headings:
            subfolder = temp_dir / heading
            subfolder.mkdir()
            doc = Document()
            doc.save(str(subfolder / f"{heading}.docx"))

        success, messages = validate_split_files(temp_dir, headings)

        assert success == True
        assert any("3/3" in msg for msg in messages)

    def test_missing_subfolders(self, temp_dir):
        """Test when some subfolders are missing."""
        headings = ["Section1", "Section2", "Section3"]

        # Only create first one
        subfolder = temp_dir / "Section1"
        subfolder.mkdir()
        doc = Document()
        doc.save(str(subfolder / "Section1.docx"))

        success, messages = validate_split_files(temp_dir, headings)

        # Success depends on ratio (>= half)
        assert any("Missing" in msg for msg in messages)

    def test_subfolder_without_docx(self, temp_dir):
        """Test subfolder exists but has no docx file."""
        headings = ["Section1"]

        subfolder = temp_dir / "Section1"
        subfolder.mkdir()
        # Don't create docx

        success, messages = validate_split_files(temp_dir, headings)

        assert any("No .docx files" in msg for msg in messages)

    def test_output_dir_not_found(self, temp_dir):
        """Test with non-existent output directory."""
        non_existent = temp_dir / "does_not_exist"

        success, messages = validate_split_files(non_existent, ["Section1"])

        assert success == False

    def test_less_than_half_present_fails(self, temp_dir):
        """Test that less than half present fails.

        Validation passes when found_count >= len(headings) // 2.
        For 5 headings, need at least 2 to pass. With only 1, it should fail.
        """
        headings = ["S1", "S2", "S3", "S4", "S5"]

        # Only create 1 of 5 (less than half of 5 which is 2)
        subfolder = temp_dir / "S1"
        subfolder.mkdir()
        doc = Document()
        doc.save(str(subfolder / "S1.docx"))

        success, messages = validate_split_files(temp_dir, headings)

        assert success == False
        assert any("Too many missing" in msg for msg in messages)

    def test_empty_headings_list(self, temp_dir):
        """Test with empty expected headings."""
        success, messages = validate_split_files(temp_dir, [])

        # Edge case - should handle gracefully


class TestValidateFormatting:
    """Tests for validate_formatting function."""

    def test_valid_formatting(self, temp_dir):
        """Test document with correct formatting."""
        doc = Document()
        para = doc.add_paragraph()
        run = para.add_run("Text with Calibri")
        run.font.name = "Calibri"

        # Add a header that should be bold
        header = doc.add_paragraph("Why This Matters")
        for run in header.runs:
            run.bold = True

        doc_path = temp_dir / "formatted.docx"
        doc.save(str(doc_path))

        success, messages = validate_formatting(doc_path)

        assert success == True

    def test_non_calibri_font_warned(self, temp_dir):
        """Test that non-Calibri fonts are warned about."""
        doc = Document()
        para = doc.add_paragraph()
        run = para.add_run("Arial text")
        run.font.name = "Arial"

        doc_path = temp_dir / "wrong_font.docx"
        doc.save(str(doc_path))

        success, messages = validate_formatting(doc_path)

        assert any("non-Calibri" in msg for msg in messages)

    def test_unbolded_header_warned(self, temp_dir):
        """Test that unbolded headers are warned about."""
        doc = Document()
        header = doc.add_paragraph("Why This Matters")
        # Don't make it bold

        doc_path = temp_dir / "unbolded.docx"
        doc.save(str(doc_path))

        success, messages = validate_formatting(doc_path)

        assert any("Unbolded" in msg for msg in messages)

    def test_file_not_found(self, temp_dir):
        """Test with non-existent file."""
        non_existent = temp_dir / "nope.docx"

        success, messages = validate_formatting(non_existent)

        assert success == False

    def test_corrupted_document(self, corrupted_docx):
        """Test with corrupted document."""
        success, messages = validate_formatting(corrupted_docx)

        assert success == False

    def test_all_checks_pass_message(self, temp_dir):
        """Test message when all checks pass."""
        doc = Document()
        para = doc.add_paragraph()
        run = para.add_run("Good text")
        run.font.name = "Calibri"

        doc_path = temp_dir / "good.docx"
        doc.save(str(doc_path))

        success, messages = validate_formatting(doc_path)

        # When no issues, should say all passed
        assert success == True


class TestValidateHyperlinks:
    """Tests for validate_hyperlinks function."""

    def test_document_with_hyperlinks(self, docx_with_hyperlinks):
        """Test document with valid hyperlinks."""
        success, messages = validate_hyperlinks(docx_with_hyperlinks)

        assert any("hyperlinks" in msg.lower() for msg in messages)

    def test_document_without_hyperlinks(self, temp_docx):
        """Test document without hyperlinks."""
        success, messages = validate_hyperlinks(temp_docx)

        assert any("0 hyperlinks" in msg or "Found 0" in msg for msg in messages)
        assert success == True

    def test_file_not_found(self, temp_dir):
        """Test with non-existent file."""
        non_existent = temp_dir / "nope.docx"

        success, messages = validate_hyperlinks(non_existent)

        assert success == False

    def test_corrupted_document(self, corrupted_docx):
        """Test with corrupted document."""
        success, messages = validate_hyperlinks(corrupted_docx)

        assert success == False

    def test_counts_multiple_hyperlinks(self, temp_dir):
        """Test counting multiple hyperlinks."""
        doc = Document()

        # Add multiple hyperlinks using helper
        from tests.conftest import add_hyperlink

        para1 = doc.add_paragraph()
        add_hyperlink(para1, "https://example.com", "Link 1")

        para2 = doc.add_paragraph()
        add_hyperlink(para2, "https://google.com", "Link 2")

        para3 = doc.add_paragraph()
        add_hyperlink(para3, "https://anthropic.com", "Link 3")

        doc_path = temp_dir / "multi_links.docx"
        doc.save(str(doc_path))

        success, messages = validate_hyperlinks(doc_path)

        assert any("3 hyperlinks" in msg for msg in messages)


class TestValidatorsIntegration:
    """Integration tests for validators."""

    def test_full_validation_workflow(self, temp_dir):
        """Test complete validation workflow."""
        # Create a document to validate
        doc = Document()
        doc.add_heading("Chief Audit Executive", level=2)
        doc.add_paragraph("Why This Matters")
        doc.add_paragraph("Content")
        doc.add_heading("General Counsel", level=2)
        doc.add_paragraph("More content")

        doc_path = temp_dir / "insights.docx"
        doc.save(str(doc_path))

        # Run all validations
        insight_success, insight_msgs = validate_insights_document(doc_path)
        format_success, format_msgs = validate_formatting(doc_path)
        hyperlink_success, hyperlink_msgs = validate_hyperlinks(doc_path)

        # All should succeed for valid document
        assert insight_success == True

    def test_validate_split_then_format(self, temp_dir):
        """Test validating split files then their formatting."""
        # Create split file structure
        headings = ["Section1", "Section2"]

        for heading in headings:
            subfolder = temp_dir / heading
            subfolder.mkdir()

            doc = Document()
            para = doc.add_paragraph()
            run = para.add_run("Content")
            run.font.name = "Calibri"
            doc.save(str(subfolder / f"{heading}.docx"))

        # Validate split
        split_success, split_msgs = validate_split_files(temp_dir, headings)
        assert split_success == True

        # Validate each file's formatting
        for heading in headings:
            file_path = temp_dir / heading / f"{heading}.docx"
            format_success, format_msgs = validate_formatting(file_path)
            assert format_success == True

    def test_all_validators_on_email_template(self, email_template_docx, temp_dir):
        """Test all validators on email template."""
        # Email validation
        email_success, email_msgs = validate_emails_document(email_template_docx)
        assert email_success == True

        # Also valid as insights document (has H2)
        insight_success, insight_msgs = validate_insights_document(email_template_docx)
        assert insight_success == True

        # Formatting validation
        format_success, format_msgs = validate_formatting(email_template_docx)
        # May have warnings but should succeed

        # Hyperlink validation
        hyperlink_success, hyperlink_msgs = validate_hyperlinks(email_template_docx)
        assert hyperlink_success == True
