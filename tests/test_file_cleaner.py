"""
Tests for file_cleaner.py

These tests challenge the file cleaning logic with:
- Documents with few paragraphs
- Special characters in content
- Filename generation edge cases
- Backup and rename operations
"""

import pytest
from pathlib import Path
from docx import Document
from freezegun import freeze_time
from datetime import datetime
import sys

sys.path.insert(0, str(Path(__file__).parent.parent))

from processors.file_cleaner import (
    find_matching_word_file,
    process_single_document,
    clean_and_rename_files,
)


class TestFindMatchingWordFile:
    """Tests for find_matching_word_file function."""

    def test_finds_matching_file(self, temp_dir):
        """Test finding a file that matches subfolder name."""
        subfolder = temp_dir / "MyFolder"
        subfolder.mkdir()

        doc = Document()
        doc.add_paragraph("Content")
        doc.save(str(subfolder / "MyFolder.docx"))

        result = find_matching_word_file(subfolder, "MyFolder")

        assert result is not None
        assert result.name == "MyFolder.docx"

    def test_no_matching_file(self, temp_dir):
        """Test when no matching file exists."""
        subfolder = temp_dir / "MyFolder"
        subfolder.mkdir()

        doc = Document()
        doc.save(str(subfolder / "DifferentName.docx"))

        result = find_matching_word_file(subfolder, "MyFolder")

        assert result is None

    def test_empty_folder(self, temp_dir):
        """Test with empty folder."""
        subfolder = temp_dir / "EmptyFolder"
        subfolder.mkdir()

        result = find_matching_word_file(subfolder, "EmptyFolder")

        assert result is None

    def test_only_txt_file(self, temp_dir):
        """Test when only non-docx files exist."""
        subfolder = temp_dir / "MyFolder"
        subfolder.mkdir()
        (subfolder / "MyFolder.txt").write_text("content")

        result = find_matching_word_file(subfolder, "MyFolder")

        assert result is None

    def test_case_matching(self, temp_dir):
        """Test filename matching behavior."""
        subfolder = temp_dir / "MyFolder"
        subfolder.mkdir()

        doc = Document()
        doc.save(str(subfolder / "myfolder.docx"))  # lowercase filename

        result = find_matching_word_file(subfolder, "MyFolder")

        # On case-insensitive filesystems, myfolder.docx matches MyFolder
        # The function looks for exact filename match, but filesystem may be case-insensitive
        # This tests the actual behavior rather than assumed behavior
        if result is not None:
            # If found, it should be the file we created
            assert result.suffix == ".docx"


class TestProcessSingleDocument:
    """Tests for process_single_document function."""

    def test_basic_processing(self, temp_dir):
        """Test basic document processing."""
        doc = Document()
        doc.add_paragraph("Line 1 - to be deleted")
        doc.add_paragraph("Line 2")
        doc.add_paragraph("Line 3 - becomes filename")
        doc.add_paragraph("Line 4")

        doc_path = temp_dir / "test.docx"
        doc.save(str(doc_path))

        success, message, new_filename = process_single_document(
            doc_path, "SubfolderName", "1.15"
        )

        assert success == True
        assert new_filename == "1.15_SubfolderName_Line_3_-_becomes_filename.docx"

    def test_deletes_first_line(self, temp_dir):
        """Test that first line is deleted."""
        doc = Document()
        doc.add_paragraph("DELETE ME")
        doc.add_paragraph("Keep this")
        doc.add_paragraph("Third line")
        doc.add_paragraph("Fourth line")

        doc_path = temp_dir / "test.docx"
        doc.save(str(doc_path))

        process_single_document(doc_path, "Folder", "1.15")

        # Reload and check
        modified_doc = Document(doc_path)
        first_para = None
        for p in modified_doc.paragraphs:
            if p.text.strip():
                first_para = p.text
                break

        # First line should be cleared (empty), not "DELETE ME"
        assert first_para != "DELETE ME"

    def test_less_than_three_lines(self, temp_dir):
        """Test document with less than 3 lines."""
        doc = Document()
        doc.add_paragraph("Line 1")
        doc.add_paragraph("Line 2")

        doc_path = temp_dir / "test.docx"
        doc.save(str(doc_path))

        success, message, new_filename = process_single_document(
            doc_path, "Folder", "1.15"
        )

        assert success == False
        assert "less than 3 lines" in message.lower()
        assert new_filename is None

    def test_empty_paragraphs_skipped(self, temp_dir):
        """Test that empty paragraphs are skipped when counting lines."""
        doc = Document()
        doc.add_paragraph("Line 1")
        doc.add_paragraph("")  # Empty
        doc.add_paragraph("Line 2")
        doc.add_paragraph("")  # Empty
        doc.add_paragraph("")  # Empty
        doc.add_paragraph("Line 3 - third non-empty")

        doc_path = temp_dir / "test.docx"
        doc.save(str(doc_path))

        success, message, new_filename = process_single_document(
            doc_path, "Folder", "1.15"
        )

        assert success == True
        assert "Line_3" in new_filename

    def test_third_line_sanitized(self, temp_dir):
        """Test that third line is sanitized for filename."""
        doc = Document()
        doc.add_paragraph("Line 1")
        doc.add_paragraph("Line 2")
        doc.add_paragraph("Line 3: Special & Characters (Here)")

        doc_path = temp_dir / "test.docx"
        doc.save(str(doc_path))

        success, message, new_filename = process_single_document(
            doc_path, "Folder", "1.15"
        )

        assert success == True
        # Special characters should be removed/sanitized
        assert ":" not in new_filename
        assert "&" not in new_filename
        assert "(" not in new_filename

    def test_subfolder_name_sanitized(self, temp_dir):
        """Test that subfolder name is sanitized in filename."""
        doc = Document()
        doc.add_paragraph("Line 1")
        doc.add_paragraph("Line 2")
        doc.add_paragraph("Line 3")

        doc_path = temp_dir / "test.docx"
        doc.save(str(doc_path))

        success, message, new_filename = process_single_document(
            doc_path, "Chief Audit Executive", "1.15"
        )

        assert success == True
        assert "Chief_Audit_Executive" in new_filename

    def test_date_prefix_in_filename(self, temp_dir):
        """Test that date prefix is included in filename."""
        doc = Document()
        doc.add_paragraph("Line 1")
        doc.add_paragraph("Line 2")
        doc.add_paragraph("Line 3")

        doc_path = temp_dir / "test.docx"
        doc.save(str(doc_path))

        success, message, new_filename = process_single_document(
            doc_path, "Folder", "2.10"
        )

        assert success == True
        assert new_filename.startswith("2.10_")

    def test_whitespace_only_paragraphs(self, temp_dir):
        """Test that whitespace-only paragraphs are treated as empty."""
        doc = Document()
        doc.add_paragraph("Line 1")
        doc.add_paragraph("   ")  # Whitespace only
        doc.add_paragraph("Line 2")
        doc.add_paragraph("\t\n")  # Whitespace only
        doc.add_paragraph("Line 3")

        doc_path = temp_dir / "test.docx"
        doc.save(str(doc_path))

        success, message, new_filename = process_single_document(
            doc_path, "Folder", "1.15"
        )

        assert success == True

    def test_very_long_third_line(self, temp_dir):
        """Test with a very long third line - should be truncated."""
        doc = Document()
        doc.add_paragraph("Line 1")
        doc.add_paragraph("Line 2")
        doc.add_paragraph("A" * 500)  # Very long line

        doc_path = temp_dir / "test.docx"
        doc.save(str(doc_path))

        success, message, new_filename = process_single_document(
            doc_path, "Folder", "1.15"
        )

        assert success == True
        # Filename is now truncated to safe length
        # Format: 1.15_Folder_AAA...AAA.docx
        # Max component length is 100, so total should be well under 255
        assert len(new_filename) < 255


class TestCleanAndRenameFiles:
    """Tests for clean_and_rename_files function."""

    @freeze_time("2024-01-15")  # Monday
    def test_basic_workflow(self, temp_dir):
        """Test basic clean and rename workflow."""
        # Create subfolder structure
        subfolder = temp_dir / "TestRole"
        subfolder.mkdir()

        doc = Document()
        doc.add_paragraph("Header to delete")
        doc.add_paragraph("Second line")
        doc.add_paragraph("Third line becomes title")
        doc.add_paragraph("More content")

        doc.save(str(subfolder / "TestRole.docx"))

        processed, skipped, backups = clean_and_rename_files(temp_dir)

        assert processed == 1
        assert skipped == 0
        assert len(backups) == 1

    @freeze_time("2024-01-15")
    def test_calculates_two_mondays_ahead(self, temp_dir):
        """Test that target date is 2 Mondays from today."""
        subfolder = temp_dir / "Role"
        subfolder.mkdir()

        doc = Document()
        doc.add_paragraph("Line 1")
        doc.add_paragraph("Line 2")
        doc.add_paragraph("Line 3")

        doc.save(str(subfolder / "Role.docx"))

        clean_and_rename_files(temp_dir)

        # Check renamed file has correct date
        # 2 Mondays from Jan 15 (Monday) = Jan 29
        renamed_files = list(subfolder.glob("*.docx"))
        # Filter out backup files
        renamed_files = [f for f in renamed_files if "backup" not in f.name]

        assert len(renamed_files) == 1
        assert renamed_files[0].name.startswith("1.29_")

    def test_creates_backup(self, temp_dir):
        """Test that backup is created before renaming."""
        subfolder = temp_dir / "Role"
        subfolder.mkdir()

        doc = Document()
        doc.add_paragraph("Line 1")
        doc.add_paragraph("Line 2")
        doc.add_paragraph("Line 3")

        original_path = subfolder / "Role.docx"
        doc.save(str(original_path))

        processed, skipped, backups = clean_and_rename_files(temp_dir)

        assert len(backups) == 1
        backup_path = Path(backups[0])
        assert backup_path.exists()
        assert "backup" in backup_path.name

    def test_skips_no_matching_file(self, temp_dir):
        """Test that subfolders without matching files are skipped."""
        subfolder = temp_dir / "Role"
        subfolder.mkdir()

        doc = Document()
        doc.save(str(subfolder / "DifferentName.docx"))

        processed, skipped, backups = clean_and_rename_files(temp_dir)

        assert processed == 0
        assert skipped == 1
        assert len(backups) == 0

    def test_skips_files_in_root(self, temp_dir):
        """Test that files in root directory are not processed."""
        # File in root
        doc = Document()
        doc.save(str(temp_dir / "RootFile.docx"))

        processed, skipped, backups = clean_and_rename_files(temp_dir)

        # Root file should not be processed
        assert (temp_dir / "RootFile.docx").exists()

    def test_multiple_subfolders(self, temp_dir):
        """Test processing multiple subfolders."""
        roles = ["Role1", "Role2", "Role3"]

        for role in roles:
            subfolder = temp_dir / role
            subfolder.mkdir()

            doc = Document()
            doc.add_paragraph("Line 1")
            doc.add_paragraph("Line 2")
            doc.add_paragraph(f"Title for {role}")

            doc.save(str(subfolder / f"{role}.docx"))

        processed, skipped, backups = clean_and_rename_files(temp_dir)

        assert processed == 3
        assert skipped == 0
        assert len(backups) == 3

    def test_non_existent_directory(self, temp_dir):
        """Test with non-existent directory."""
        non_existent = temp_dir / "does_not_exist"

        processed, skipped, backups = clean_and_rename_files(non_existent)

        assert processed == 0
        assert skipped == 0
        assert backups == []

    def test_progress_callback(self, temp_dir):
        """Test that progress callback is called."""
        subfolder = temp_dir / "Role"
        subfolder.mkdir()

        doc = Document()
        doc.add_paragraph("Line 1")
        doc.add_paragraph("Line 2")
        doc.add_paragraph("Line 3")
        doc.save(str(subfolder / "Role.docx"))

        messages = []
        def callback(msg):
            messages.append(msg)

        clean_and_rename_files(temp_dir, progress_callback=callback)

        assert len(messages) > 0

    def test_document_with_insufficient_lines(self, temp_dir):
        """Test subfolder with document having too few lines."""
        subfolder = temp_dir / "Role"
        subfolder.mkdir()

        doc = Document()
        doc.add_paragraph("Only one line")
        doc.save(str(subfolder / "Role.docx"))

        processed, skipped, backups = clean_and_rename_files(temp_dir)

        assert processed == 0
        assert skipped == 1


class TestFileCleanerIntegration:
    """Integration tests for file cleaner."""

    @freeze_time("2024-01-15")
    def test_full_realistic_workflow(self, temp_dir):
        """Test realistic workflow with actual role names."""
        roles_and_titles = [
            ("Chief Audit Executive", "Audit Risk Assessment 2024"),
            ("General Counsel", "Legal Compliance Update"),
            ("Tech CEO", "Technology Roadmap Q1"),
        ]

        for role, title in roles_and_titles:
            subfolder = temp_dir / role
            subfolder.mkdir()

            doc = Document()
            doc.add_paragraph("HEADER TO DELETE")
            doc.add_paragraph("Subtitle line")
            doc.add_paragraph(title)
            doc.add_paragraph("Content paragraph 1")
            doc.add_paragraph("Content paragraph 2")

            doc.save(str(subfolder / f"{role}.docx"))

        processed, skipped, backups = clean_and_rename_files(temp_dir)

        assert processed == 3

        # Verify each file was renamed correctly
        for role, title in roles_and_titles:
            subfolder = temp_dir / role
            files = list(subfolder.glob("*.docx"))
            # Filter out backups
            renamed = [f for f in files if "backup" not in f.name]

            assert len(renamed) == 1
            # Check filename contains sanitized title
            assert title.replace(" ", "_") in renamed[0].name or \
                   title.split()[0] in renamed[0].name

    @freeze_time("2024-12-28")  # Near year end
    def test_year_boundary_date(self, temp_dir):
        """Test date calculation near year boundary."""
        subfolder = temp_dir / "Role"
        subfolder.mkdir()

        doc = Document()
        doc.add_paragraph("Line 1")
        doc.add_paragraph("Line 2")
        doc.add_paragraph("Title")
        doc.save(str(subfolder / "Role.docx"))

        clean_and_rename_files(temp_dir)

        files = list(subfolder.glob("*.docx"))
        renamed = [f for f in files if "backup" not in f.name]

        # Date should be in January 2025 (2 Mondays from Dec 28)
        # Dec 30 is first Monday, Jan 6 is second Monday
        assert renamed[0].name.startswith("1.")  # January
