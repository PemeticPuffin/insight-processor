"""
Shared pytest fixtures for Insight Processor tests.
"""

import pytest
import tempfile
import shutil
from pathlib import Path
from datetime import datetime, timedelta
from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


@pytest.fixture
def temp_dir():
    """Create a temporary directory for test files."""
    temp_path = Path(tempfile.mkdtemp())
    yield temp_path
    shutil.rmtree(temp_path, ignore_errors=True)


@pytest.fixture
def temp_docx(temp_dir):
    """Create a simple temporary Word document."""
    doc_path = temp_dir / "test_document.docx"
    doc = Document()
    doc.add_paragraph("Test paragraph")
    doc.save(str(doc_path))
    return doc_path


@pytest.fixture
def complex_docx(temp_dir):
    """Create a Word document with multiple paragraphs and styles."""
    doc_path = temp_dir / "complex_document.docx"
    doc = Document()

    # Add various content
    doc.add_heading("Heading 1 Title", level=1)
    doc.add_heading("Section One", level=2)
    doc.add_paragraph("First paragraph content.")
    doc.add_paragraph("Second paragraph content.")

    doc.add_heading("Section Two", level=2)
    doc.add_paragraph("Third paragraph content.")

    doc.add_heading("Section Three", level=2)
    doc.add_paragraph("Fourth paragraph content.")
    doc.add_paragraph("Fifth paragraph content.")
    doc.add_paragraph("Sixth paragraph content.")

    doc.save(str(doc_path))
    return doc_path


@pytest.fixture
def docx_with_headers(temp_dir):
    """Create a Word document with specific headers matching config."""
    doc_path = temp_dir / "headers_document.docx"
    doc = Document()

    # Add title
    title = doc.add_paragraph("Title of Insight: Test Insight")

    # Add headers that match HEADERS_TO_BOLD
    doc.add_paragraph("Why This Matters")
    doc.add_paragraph("This is important content about why it matters.")
    doc.add_paragraph("Here is more context for the reader.")

    doc.add_paragraph("Implications")
    doc.add_paragraph("First implication of the insight.")
    doc.add_paragraph("Second implication to consider.")

    doc.add_paragraph("Why Gartner")
    doc.add_paragraph("Gartner provides expertise in this area.")

    doc.add_paragraph("Probing Questions")
    doc.add_paragraph("What challenges are you facing?")
    doc.add_paragraph("How does your team handle this?")

    doc.add_paragraph("AskGartner Prompts to Recommend")
    doc.add_paragraph("Ask about market trends.")
    doc.add_paragraph("Ask about best practices.")

    doc.save(str(doc_path))
    return doc_path


@pytest.fixture
def docx_with_hyperlinks(temp_dir):
    """Create a Word document with hyperlinks."""
    doc_path = temp_dir / "hyperlinks_document.docx"
    doc = Document()

    # Add paragraph with hyperlink
    para = doc.add_paragraph()
    add_hyperlink(para, "https://www.example.com", "Example Link")

    para2 = doc.add_paragraph("Normal text ")
    add_hyperlink(para2, "https://www.google.com", "Google")

    doc.save(str(doc_path))
    return doc_path


@pytest.fixture
def docx_with_lists(temp_dir):
    """Create a Word document with numbered and bulleted lists."""
    doc_path = temp_dir / "lists_document.docx"
    doc = Document()

    doc.add_paragraph("Introduction paragraph.")

    # Numbered list items (as plain text)
    doc.add_paragraph("1. First numbered item")
    doc.add_paragraph("2. Second numbered item")
    doc.add_paragraph("3. Third numbered item")

    doc.add_paragraph("Transition paragraph.")

    # Dash list items
    doc.add_paragraph("- First dash item")
    doc.add_paragraph("- Second dash item")

    # Letter list items
    doc.add_paragraph("a. First letter item")
    doc.add_paragraph("b. Second letter item")

    doc.save(str(doc_path))
    return doc_path


@pytest.fixture
def email_template_docx(temp_dir):
    """Create a Word document with email template structure."""
    doc_path = temp_dir / "email_templates.docx"
    doc = Document()

    # Job role 1
    doc.add_heading("Chief Audit Executive", level=2)
    doc.add_heading("SALES BUSINESS DEVELOPER PROSPECT EMAIL", level=3)
    doc.add_paragraph("BD email content for CAE.")
    doc.add_heading("SALES ACCOUNT EXECUTIVE CLIENT EMAIL", level=3)
    doc.add_paragraph("AE email content for CAE.")
    doc.add_heading("SERVICE EXECUTIVE PARTNER CLIENT EMAIL", level=3)
    doc.add_paragraph("EP email content for CAE.")

    # Job role 2
    doc.add_heading("General Counsel", level=2)
    doc.add_heading("SALES BUSINESS DEVELOPER PROSPECT EMAIL", level=3)
    doc.add_paragraph("BD email content for GC.")
    doc.add_heading("SALES ACCOUNT EXECUTIVE CLIENT EMAIL", level=3)
    doc.add_paragraph("AE email content for GC.")

    doc.save(str(doc_path))
    return doc_path


@pytest.fixture
def monday_dates():
    """Return a list of known Monday dates for testing."""
    # Known Mondays
    return [
        datetime(2024, 1, 1),   # Monday
        datetime(2024, 1, 8),   # Monday
        datetime(2024, 1, 15),  # Monday
        datetime(2024, 2, 5),   # Monday
        datetime(2024, 12, 30), # Monday
    ]


@pytest.fixture
def non_monday_dates():
    """Return dates that are NOT Mondays with their expected next Monday."""
    return [
        (datetime(2024, 1, 2), datetime(2024, 1, 8)),   # Tuesday -> next Monday
        (datetime(2024, 1, 3), datetime(2024, 1, 8)),   # Wednesday -> next Monday
        (datetime(2024, 1, 4), datetime(2024, 1, 8)),   # Thursday -> next Monday
        (datetime(2024, 1, 5), datetime(2024, 1, 8)),   # Friday -> next Monday
        (datetime(2024, 1, 6), datetime(2024, 1, 8)),   # Saturday -> next Monday
        (datetime(2024, 1, 7), datetime(2024, 1, 8)),   # Sunday -> next Monday
    ]


def add_hyperlink(paragraph, url, text, color=None, underline=True):
    """
    Add a hyperlink to a paragraph.

    Args:
        paragraph: The paragraph to add the hyperlink to
        url: The URL for the hyperlink
        text: The display text for the hyperlink
        color: Optional color for the hyperlink text
        underline: Whether to underline the hyperlink text

    Returns:
        The hyperlink element
    """
    part = paragraph.part
    r_id = part.relate_to(url, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink", is_external=True)

    hyperlink = OxmlElement('w:hyperlink')
    hyperlink.set(qn('r:id'), r_id)

    new_run = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')

    if underline:
        u = OxmlElement('w:u')
        u.set(qn('w:val'), 'single')
        rPr.append(u)

    if color:
        c = OxmlElement('w:color')
        c.set(qn('w:val'), color)
        rPr.append(c)

    new_run.append(rPr)

    t = OxmlElement('w:t')
    t.text = text
    new_run.append(t)

    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)

    return hyperlink


@pytest.fixture
def corrupted_docx(temp_dir):
    """Create a corrupted/invalid docx file."""
    doc_path = temp_dir / "corrupted.docx"
    # Write invalid content (not a valid zip/docx)
    with open(doc_path, 'w') as f:
        f.write("This is not a valid docx file")
    return doc_path


@pytest.fixture
def empty_docx(temp_dir):
    """Create an empty Word document."""
    doc_path = temp_dir / "empty.docx"
    doc = Document()
    doc.save(str(doc_path))
    return doc_path


@pytest.fixture
def nested_dir_structure(temp_dir):
    """Create a nested directory structure with Word files."""
    # Create subdirectories
    sub1 = temp_dir / "subfolder1"
    sub2 = temp_dir / "subfolder2"
    nested = sub1 / "nested"

    sub1.mkdir()
    sub2.mkdir()
    nested.mkdir()

    # Create docx files
    for i, folder in enumerate([temp_dir, sub1, sub2, nested], 1):
        doc = Document()
        doc.add_paragraph(f"Document in {folder.name}")
        doc.save(str(folder / f"doc{i}.docx"))

    # Create a temp file that should be ignored
    with open(sub1 / "~$tempfile.docx", 'w') as f:
        f.write("temp")

    # Create non-docx files
    (sub2 / "readme.txt").write_text("readme content")
    (sub2 / "data.xlsx").write_text("excel content")

    return temp_dir
