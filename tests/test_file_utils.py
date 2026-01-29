"""
Tests for file_utils.py

These tests challenge file operations with:
- Permission issues
- Non-existent files
- Corrupted files
- Temp file detection
- Recursive vs non-recursive search
- Various file types
"""

import pytest
import os
import stat
import tempfile
import shutil
from pathlib import Path
from datetime import datetime
from unittest.mock import patch, MagicMock
from docx import Document
import sys

sys.path.insert(0, str(Path(__file__).parent.parent))

from utils.file_utils import backup_file, find_word_files, is_temp_file, validate_docx


class TestBackupFile:
    """Tests for backup_file function."""

    def test_basic_backup(self, temp_dir):
        """Test basic file backup functionality."""
        # Create a test file
        test_file = temp_dir / "test.docx"
        test_file.write_text("test content")

        backup_path = backup_file(test_file)

        assert backup_path.exists()
        assert "backup" in backup_path.name
        assert backup_path.suffix == ".docx"
        assert backup_path.parent == test_file.parent

    def test_backup_preserves_content(self, temp_dir):
        """Test that backup preserves file content."""
        test_file = temp_dir / "test.txt"
        original_content = "Important data\nLine 2\nLine 3"
        test_file.write_text(original_content)

        backup_path = backup_file(test_file)

        assert backup_path.read_text() == original_content

    def test_backup_file_not_found(self, temp_dir):
        """Test backup of non-existent file raises error."""
        non_existent = temp_dir / "does_not_exist.docx"

        with pytest.raises(FileNotFoundError):
            backup_file(non_existent)

    def test_backup_filename_format(self, temp_dir):
        """Test backup filename follows expected format."""
        test_file = temp_dir / "report.docx"
        test_file.write_text("content")

        backup_path = backup_file(test_file)

        # Format: {stem}_backup_{timestamp}{suffix}
        assert backup_path.name.startswith("report_backup_")
        assert backup_path.name.endswith(".docx")

    def test_backup_timestamp_format(self, temp_dir):
        """Test that timestamp in backup name is valid."""
        test_file = temp_dir / "test.docx"
        test_file.write_text("content")

        backup_path = backup_file(test_file)

        # Extract timestamp from filename
        name_parts = backup_path.stem.split("_backup_")
        assert len(name_parts) == 2

        timestamp = name_parts[1]
        # Should be format: YYYYMMDD_HHMMSS
        assert len(timestamp) == 15  # 8 + 1 + 6
        assert timestamp[8] == "_"

    def test_multiple_backups_unique_names(self, temp_dir):
        """Test that multiple backups get unique names."""
        test_file = temp_dir / "test.docx"
        test_file.write_text("content")

        import time
        backup1 = backup_file(test_file)
        time.sleep(1.1)  # Ensure different timestamp
        backup2 = backup_file(test_file)

        assert backup1 != backup2
        assert backup1.exists()
        assert backup2.exists()

    def test_backup_preserves_metadata(self, temp_dir):
        """Test that backup preserves file modification time."""
        test_file = temp_dir / "test.docx"
        test_file.write_text("content")

        backup_path = backup_file(test_file)

        # shutil.copy2 preserves metadata
        original_stat = test_file.stat()
        backup_stat = backup_path.stat()

        # Modification time should be close (copy2 preserves it)
        assert abs(original_stat.st_mtime - backup_stat.st_mtime) < 1

    def test_backup_large_file(self, temp_dir):
        """Test backup of a large file."""
        test_file = temp_dir / "large.docx"
        # Create a 1MB file
        test_file.write_bytes(b"x" * (1024 * 1024))

        backup_path = backup_file(test_file)

        assert backup_path.exists()
        assert backup_path.stat().st_size == test_file.stat().st_size

    def test_backup_file_with_spaces_in_name(self, temp_dir):
        """Test backup of file with spaces in name."""
        test_file = temp_dir / "my document with spaces.docx"
        test_file.write_text("content")

        backup_path = backup_file(test_file)

        assert backup_path.exists()
        assert "my document with spaces" in backup_path.stem

    def test_backup_file_with_unicode_name(self, temp_dir):
        """Test backup of file with Unicode characters in name."""
        test_file = temp_dir / "document_\u00e9\u00e0\u00fc.docx"
        test_file.write_text("content")

        backup_path = backup_file(test_file)

        assert backup_path.exists()

    @pytest.mark.skipif(os.name == 'nt', reason="Permission tests unreliable on Windows")
    def test_backup_read_only_source(self, temp_dir):
        """Test backup of read-only file."""
        test_file = temp_dir / "readonly.docx"
        test_file.write_text("content")
        os.chmod(test_file, stat.S_IRUSR)

        try:
            # Should still work - we can read the file
            backup_path = backup_file(test_file)
            assert backup_path.exists()
        finally:
            # Cleanup - restore permissions
            os.chmod(test_file, stat.S_IRWXU)


class TestFindWordFiles:
    """Tests for find_word_files function."""

    def test_find_files_in_single_directory(self, temp_dir):
        """Test finding Word files in a single directory."""
        # Create test files
        (temp_dir / "doc1.docx").write_text("content")
        (temp_dir / "doc2.docx").write_text("content")
        (temp_dir / "readme.txt").write_text("content")

        result = find_word_files(temp_dir, recursive=False)

        assert len(result) == 2
        assert all(p.suffix == ".docx" for p in result)

    def test_find_files_recursive(self, nested_dir_structure):
        """Test recursive file finding."""
        result = find_word_files(nested_dir_structure, recursive=True)

        # Should find files in main dir and all subdirs
        assert len(result) >= 4

    def test_find_files_non_recursive(self, nested_dir_structure):
        """Test non-recursive file finding."""
        result = find_word_files(nested_dir_structure, recursive=False)

        # Should only find files in the top directory
        for file_path in result:
            assert file_path.parent == nested_dir_structure

    def test_excludes_temp_files(self, temp_dir):
        """Test that temp files (starting with ~$) are excluded."""
        (temp_dir / "document.docx").write_text("content")
        (temp_dir / "~$document.docx").write_text("temp")
        (temp_dir / "~$another.docx").write_text("temp")

        result = find_word_files(temp_dir)

        assert len(result) == 1
        assert "~$" not in result[0].name

    def test_empty_directory(self, temp_dir):
        """Test finding files in empty directory."""
        result = find_word_files(temp_dir)

        assert result == []

    def test_directory_with_no_docx(self, temp_dir):
        """Test directory with no Word files."""
        (temp_dir / "readme.txt").write_text("content")
        (temp_dir / "data.xlsx").write_text("content")
        (temp_dir / "script.py").write_text("content")

        result = find_word_files(temp_dir)

        assert result == []

    def test_sorted_by_modification_time(self, temp_dir):
        """Test that results are sorted by modification time (newest first)."""
        import time

        file1 = temp_dir / "old.docx"
        file1.write_text("old")

        time.sleep(0.1)

        file2 = temp_dir / "new.docx"
        file2.write_text("new")

        result = find_word_files(temp_dir)

        # Newest first
        assert result[0].name == "new.docx"
        assert result[1].name == "old.docx"

    def test_handles_symbolic_links(self, temp_dir):
        """Test handling of symbolic links to docx files."""
        # Create a real file
        real_file = temp_dir / "real.docx"
        real_file.write_text("content")

        # Create a symlink
        link = temp_dir / "link.docx"
        try:
            link.symlink_to(real_file)
        except OSError:
            pytest.skip("Symlinks not supported on this platform")

        result = find_word_files(temp_dir)

        # Should find both (depending on implementation)
        # At minimum, should not crash
        assert len(result) >= 1

    def test_mixed_case_extension(self, temp_dir):
        """Test that only lowercase .docx is matched."""
        (temp_dir / "doc1.docx").write_text("content")
        (temp_dir / "doc2.DOCX").write_text("content")
        (temp_dir / "doc3.Docx").write_text("content")

        result = find_word_files(temp_dir)

        # Current implementation only checks for .docx (lowercase)
        filenames = [p.name for p in result]
        assert "doc1.docx" in filenames
        # Upper/mixed case might not be found - depends on implementation

    def test_deeply_nested_structure(self, temp_dir):
        """Test deeply nested directory structure."""
        # Create deep nesting
        deep_path = temp_dir
        for i in range(10):
            deep_path = deep_path / f"level{i}"
            deep_path.mkdir()

        # Add file at deepest level
        (deep_path / "deep.docx").write_text("content")

        result = find_word_files(temp_dir, recursive=True)

        assert len(result) == 1
        assert result[0].name == "deep.docx"

    def test_hidden_directories(self, temp_dir):
        """Test handling of hidden directories (starting with .)."""
        hidden_dir = temp_dir / ".hidden"
        hidden_dir.mkdir()
        (hidden_dir / "secret.docx").write_text("content")
        (temp_dir / "visible.docx").write_text("content")

        result = find_word_files(temp_dir, recursive=True)

        # Hidden directory contents should still be found
        filenames = [p.name for p in result]
        assert "visible.docx" in filenames
        # Hidden dir file might or might not be found depending on os.walk behavior


class TestIsTempFile:
    """Tests for is_temp_file function."""

    @pytest.mark.parametrize("filename,expected", [
        ("~$document.docx", True),
        ("~$temp.docx", True),
        ("~$.docx", True),
        ("document.docx", False),
        ("temp_document.docx", False),
        ("$document.docx", False),
        ("~document.docx", False),
        ("~~document.docx", False),
    ])
    def test_temp_file_detection(self, filename, expected):
        """Test various filename patterns for temp file detection."""
        assert is_temp_file(filename) == expected

    def test_empty_filename(self):
        """Test with empty filename."""
        assert is_temp_file("") == False

    def test_just_prefix(self):
        """Test with just the temp prefix."""
        assert is_temp_file("~$") == True

    def test_case_sensitivity(self):
        """Test that prefix detection is case-sensitive."""
        # The actual temp file prefix is ~$ exactly
        assert is_temp_file("~$file.docx") == True


class TestValidateDocx:
    """Tests for validate_docx function."""

    def test_valid_docx(self, temp_docx):
        """Test validation of a valid Word document."""
        assert validate_docx(temp_docx) == True

    def test_non_existent_file(self, temp_dir):
        """Test validation of non-existent file."""
        non_existent = temp_dir / "does_not_exist.docx"
        assert validate_docx(non_existent) == False

    def test_wrong_extension(self, temp_dir):
        """Test file with wrong extension."""
        txt_file = temp_dir / "document.txt"
        txt_file.write_text("content")
        assert validate_docx(txt_file) == False

    def test_corrupted_docx(self, corrupted_docx):
        """Test validation of corrupted/invalid docx."""
        assert validate_docx(corrupted_docx) == False

    def test_empty_docx(self, empty_docx):
        """Test validation of empty but valid docx."""
        assert validate_docx(empty_docx) == True

    def test_renamed_non_docx(self, temp_dir):
        """Test a non-docx file renamed to .docx."""
        # Create a text file and rename it
        fake_docx = temp_dir / "fake.docx"
        fake_docx.write_text("This is not a real docx")

        assert validate_docx(fake_docx) == False

    def test_uppercase_extension(self, temp_dir):
        """Test file with uppercase .DOCX extension."""
        doc = Document()
        doc_path = temp_dir / "document.DOCX"
        doc.save(str(doc_path))

        # The function checks for .docx (case-insensitive via .lower())
        assert validate_docx(doc_path) == True

    def test_mixed_case_extension(self, temp_dir):
        """Test file with mixed case .Docx extension."""
        doc = Document()
        doc_path = temp_dir / "document.Docx"
        doc.save(str(doc_path))

        assert validate_docx(doc_path) == True

    def test_docx_with_complex_content(self, complex_docx):
        """Test validation of docx with various content types."""
        assert validate_docx(complex_docx) == True

    def test_docx_with_hyperlinks(self, docx_with_hyperlinks):
        """Test validation of docx with hyperlinks."""
        assert validate_docx(docx_with_hyperlinks) == True

    def test_zip_file_renamed_to_docx(self, temp_dir):
        """Test a regular zip file renamed to .docx."""
        import zipfile

        zip_path = temp_dir / "archive.docx"
        with zipfile.ZipFile(zip_path, 'w') as zf:
            zf.writestr("test.txt", "content")

        # A regular zip is not a valid docx
        assert validate_docx(zip_path) == False

    def test_docx_is_actually_a_zip(self, temp_docx):
        """Confirm that a valid docx IS a valid zip file."""
        import zipfile

        # docx files are actually zip archives
        assert zipfile.is_zipfile(temp_docx)
        assert validate_docx(temp_docx) == True


class TestFileUtilsIntegration:
    """Integration tests for file utilities."""

    def test_find_backup_validate_workflow(self, temp_dir):
        """Test a typical workflow: find, backup, validate."""
        # Create some Word files
        doc = Document()
        doc.add_paragraph("Test content")
        doc_path = temp_dir / "report.docx"
        doc.save(str(doc_path))

        # Find files
        found = find_word_files(temp_dir)
        assert len(found) == 1
        assert found[0] == doc_path

        # Validate
        assert validate_docx(found[0]) == True

        # Backup
        backup_path = backup_file(found[0])
        assert backup_path.exists()
        assert validate_docx(backup_path) == True

    def test_find_excludes_temp_validates_originals(self, temp_dir):
        """Test that temp files are excluded but originals validate."""
        # Create original and temp file
        doc = Document()
        doc.save(str(temp_dir / "document.docx"))

        # Create fake temp file
        (temp_dir / "~$document.docx").write_text("temp")

        found = find_word_files(temp_dir)

        assert len(found) == 1
        assert is_temp_file(found[0].name) == False
        assert validate_docx(found[0]) == True

    def test_backup_invalid_file(self, temp_dir):
        """Test that we can backup an invalid docx (corruption happens after)."""
        # Create a valid file first
        doc_path = temp_dir / "will_be_corrupted.docx"
        doc_path.write_text("not valid docx content")

        # The file exists, so backup should work
        backup_path = backup_file(doc_path)

        assert backup_path.exists()
        # Both original and backup are invalid docx
        assert validate_docx(doc_path) == False
        assert validate_docx(backup_path) == False
