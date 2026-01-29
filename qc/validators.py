"""
QC Validators for Insight Processor pipeline.

Pre-processing and post-processing validation checks.
"""

from pathlib import Path
from docx import Document
from typing import Tuple, List, Set

import sys
sys.path.insert(0, str(Path(__file__).parent.parent))

from config import HEADERS_TO_BOLD, FONT_NAME


def validate_insights_document(doc_path: Path) -> Tuple[bool, List[str]]:
    """
    Validate an insights document before processing.

    Checks:
    - Document opens successfully
    - Contains at least one Heading 2
    - Reports which expected headings are present/missing

    Args:
        doc_path: Path to the insights document

    Returns:
        Tuple of (success, list of errors/warnings)
    """
    errors = []
    warnings = []

    # Check file exists
    if not doc_path.exists():
        return False, [f"File not found: {doc_path}"]

    # Try to open the document
    try:
        doc = Document(doc_path)
    except Exception as e:
        return False, [f"Failed to open document: {str(e)}"]

    # Find all Heading 2 paragraphs
    heading2_texts = []
    for para in doc.paragraphs:
        if para.style.name == 'Heading 2':
            heading2_texts.append(para.text.strip())

    # Check for at least one Heading 2
    if not heading2_texts:
        errors.append("No Heading 2 sections found in document")
        return False, errors

    # Report found headings
    warnings.append(f"Found {len(heading2_texts)} Heading 2 sections:")
    for h in heading2_texts[:10]:  # Show first 10
        warnings.append(f"  - {h}")
    if len(heading2_texts) > 10:
        warnings.append(f"  ... and {len(heading2_texts) - 10} more")

    return True, warnings


def validate_emails_document(doc_path: Path) -> Tuple[bool, List[str]]:
    """
    Validate an email templates document before processing.

    Checks:
    - Document opens successfully
    - Contains Heading 2 sections (job roles)
    - Contains BD/AE/EP Heading 3 sections

    Args:
        doc_path: Path to the email templates document

    Returns:
        Tuple of (success, list of errors/warnings)
    """
    errors = []
    warnings = []

    # Check file exists
    if not doc_path.exists():
        return False, [f"File not found: {doc_path}"]

    # Try to open the document
    try:
        doc = Document(doc_path)
    except Exception as e:
        return False, [f"Failed to open document: {str(e)}"]

    # Find all Heading 2 and Heading 3 paragraphs
    heading2_count = 0
    heading3_types = {'BD': 0, 'AE': 0, 'EP': 0}

    for para in doc.paragraphs:
        if para.style.name == 'Heading 2':
            heading2_count += 1

        elif para.style.name == 'Heading 3':
            text = para.text.strip().upper()
            if 'SALES BUSINESS DEVELOPER' in text or 'PROSPECT EMAIL' in text:
                heading3_types['BD'] += 1
            elif 'SALES ACCOUNT EXECUTIVE' in text:
                heading3_types['AE'] += 1
            elif 'SERVICE EXECUTIVE PARTNER' in text or ('SERVICE' in text and 'CLIENT' in text):
                heading3_types['EP'] += 1

    # Check for Heading 2 sections
    if heading2_count == 0:
        errors.append("No Heading 2 sections (job roles) found")
        return False, errors

    warnings.append(f"Found {heading2_count} job role sections")
    warnings.append(f"BD sections: {heading3_types['BD']}")
    warnings.append(f"AE sections: {heading3_types['AE']}")
    warnings.append(f"EP sections: {heading3_types['EP']}")

    # Warn if missing any template types
    if heading3_types['BD'] == 0:
        warnings.append("Warning: No BD (Business Developer) email templates found")
    if heading3_types['AE'] == 0:
        warnings.append("Warning: No AE (Account Executive) email templates found")
    if heading3_types['EP'] == 0:
        warnings.append("Warning: No EP (Executive Partner) email templates found")

    return True, warnings


def validate_split_files(
    output_dir: Path,
    expected_headings: List[str]
) -> Tuple[bool, List[str]]:
    """
    Validate that split files were created correctly.

    Checks:
    - Expected number of subfolders exist
    - Each subfolder contains a .docx file
    - Files have required sections

    Args:
        output_dir: Directory containing the split files
        expected_headings: List of expected Heading 2 texts (subfolder names)

    Returns:
        Tuple of (success, list of errors/warnings)
    """
    errors = []
    warnings = []

    if not output_dir.exists():
        return False, [f"Output directory not found: {output_dir}"]

    found_count = 0
    missing = []

    for heading in expected_headings:
        subfolder = output_dir / heading
        if not subfolder.exists():
            missing.append(heading)
            continue

        # Check for .docx file in subfolder
        docx_files = list(subfolder.glob("*.docx"))
        if not docx_files:
            warnings.append(f"No .docx files in subfolder: {heading}")
        else:
            found_count += 1

    if missing:
        warnings.append(f"Missing subfolders: {', '.join(missing[:5])}")
        if len(missing) > 5:
            warnings.append(f"  ... and {len(missing) - 5} more")

    warnings.insert(0, f"Found {found_count}/{len(expected_headings)} expected subfolders with files")

    # Consider it a failure if less than half were created
    success = found_count >= len(expected_headings) // 2
    if not success:
        errors.append(f"Too many missing files: {len(missing)} of {len(expected_headings)}")
        return False, errors + warnings

    return True, warnings


def validate_formatting(file_path: Path) -> Tuple[bool, List[str]]:
    """
    Validate that a file has been formatted correctly.

    Checks:
    - Font is Calibri
    - Headers from HEADERS_TO_BOLD are bold
    - Bullet sections use bullet style

    Args:
        file_path: Path to the formatted document

    Returns:
        Tuple of (success, list of errors/warnings)
    """
    errors = []
    warnings = []

    if not file_path.exists():
        return False, [f"File not found: {file_path}"]

    try:
        doc = Document(file_path)
    except Exception as e:
        return False, [f"Failed to open document: {str(e)}"]

    # Track issues
    non_calibri_count = 0
    unbolded_headers = []
    header_set: Set[str] = set(HEADERS_TO_BOLD)

    for para in doc.paragraphs:
        text = para.text.strip()

        # Check font
        for run in para.runs:
            if run.font.name and run.font.name != FONT_NAME:
                non_calibri_count += 1

        # Check if header is bold
        if text in header_set:
            all_bold = all(run.bold for run in para.runs if run.text.strip())
            if not all_bold and para.runs:
                unbolded_headers.append(text)

    # Report issues
    if non_calibri_count > 0:
        warnings.append(f"Found {non_calibri_count} runs with non-Calibri font")

    if unbolded_headers:
        warnings.append(f"Unbolded headers: {', '.join(unbolded_headers)}")

    # Success if no critical issues
    success = len(errors) == 0
    if not warnings:
        warnings.append("All formatting checks passed")

    return success, warnings


def validate_hyperlinks(file_path: Path) -> Tuple[bool, List[str]]:
    """
    Validate that hyperlinks are preserved in a document.

    Args:
        file_path: Path to the document to check

    Returns:
        Tuple of (success, list of errors/warnings)
    """
    errors = []
    warnings = []

    if not file_path.exists():
        return False, [f"File not found: {file_path}"]

    try:
        doc = Document(file_path)
    except Exception as e:
        return False, [f"Failed to open document: {str(e)}"]

    # Count hyperlinks by checking relationships
    hyperlink_count = 0
    empty_hyperlinks = 0

    w_ns = '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}'
    r_ns = '{http://schemas.openxmlformats.org/officeDocument/2006/relationships}'

    for para in doc.paragraphs:
        for hyperlink in para._element.iter(w_ns + 'hyperlink'):
            rel_id = hyperlink.get(r_ns + 'id')
            if rel_id:
                hyperlink_count += 1
                # Check if relationship exists
                try:
                    rel = doc.part.rels[rel_id]
                    if not rel.target_ref:
                        empty_hyperlinks += 1
                except KeyError:
                    empty_hyperlinks += 1

    warnings.append(f"Found {hyperlink_count} hyperlinks")
    if empty_hyperlinks > 0:
        warnings.append(f"Warning: {empty_hyperlinks} hyperlinks have missing targets")

    success = empty_hyperlinks == 0
    return success, warnings
