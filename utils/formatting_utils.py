"""
Formatting utility functions for fonts, bullets, and headers.
"""

import re
from docx.shared import Pt, Inches
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from config import FONT_NAME, FONT_SIZE_PT

# Bullet character for manual bullet formatting
BULLET_CHAR = '\u2022'  # •


def set_font_calibri_11pt(paragraph):
    """
    Set all runs in a paragraph to Calibri 11pt font.

    Args:
        paragraph: The paragraph to format
    """
    for run in paragraph.runs:
        run.font.name = FONT_NAME
        run.font.size = Pt(FONT_SIZE_PT)


def bold_paragraph(paragraph):
    """
    Make all runs in a paragraph bold.

    Args:
        paragraph: The paragraph to bold
    """
    for run in paragraph.runs:
        run.bold = True


def clear_all_list_formatting(paragraph):
    """
    Aggressively remove ALL list-related formatting from a paragraph.

    This clears numPr and any other list-related elements from the paragraph XML.
    Also removes the pStyle if it references a list style.

    Args:
        paragraph: The paragraph to clear list formatting from
    """
    p_element = paragraph._element
    pPr = p_element.find(qn('w:pPr'))

    if pPr is not None:
        # Remove numPr (numbering properties)
        for num_element in pPr.findall(qn('w:numPr')):
            pPr.remove(num_element)

        # Remove any list-related style references
        for pStyle in pPr.findall(qn('w:pStyle')):
            style_val = pStyle.get(qn('w:val'), '')
            if 'List' in style_val or 'Bullet' in style_val or 'Number' in style_val:
                pPr.remove(pStyle)

    # Also check for numPr directly under the paragraph element (rare but possible)
    for num_element in p_element.findall(qn('w:numPr')):
        p_element.remove(num_element)


def apply_bullet_style(paragraph):
    """
    Apply bullet point formatting to a paragraph.

    Extracts clean text (stripping any existing bullet chars, numbers, or dashes),
    clears all paragraph content and list formatting, then rebuilds with a dot bullet.

    Args:
        paragraph: The paragraph to convert to a bullet
    """
    # Extract clean text (strips bullet chars, numbered markers, dashes)
    text = _extract_clean_text(paragraph.text)

    # Clear all list formatting
    clear_all_list_formatting(paragraph)

    # Set to Normal style
    paragraph.style = 'Normal'

    # Clear again after style change (style can re-add formatting)
    clear_all_list_formatting(paragraph)

    # Rebuild the paragraph content
    paragraph.clear()

    # Add fresh run with bullet character, tab, and text
    run = paragraph.add_run(BULLET_CHAR + '\t' + text)
    run.font.name = FONT_NAME
    run.font.size = Pt(FONT_SIZE_PT)

    # Set indentation for bullet appearance
    paragraph.paragraph_format.left_indent = Inches(0.5)
    paragraph.paragraph_format.first_line_indent = Inches(-0.25)


def _extract_clean_text(text: str) -> str:
    """
    Extract clean text from a paragraph, stripping bullet chars and list markers.

    Args:
        text: The raw paragraph text

    Returns:
        Clean text with markers removed
    """
    text = text.strip()

    # Strip bullet char
    if text.startswith(BULLET_CHAR):
        text = text[1:].lstrip('\t ')

    # Strip numbered list markers (1. / 1) / (1) / a. / a) / (a))
    text = re.sub(r'^\d+[\.\)]\s*', '', text)
    text = re.sub(r'^[a-zA-Z][\.\)]\s*', '', text)
    text = re.sub(r'^\(\d+\)\s*', '', text)
    text = re.sub(r'^\([a-zA-Z]\)\s*', '', text)

    # Strip dash/hyphen markers
    for prefix in ['- ', '– ', '— ']:
        if text.startswith(prefix):
            text = text[len(prefix):]
            break

    return text.strip()


def apply_bullet_with_hyperlink(paragraph, base_url, doc):
    """
    Apply bullet point formatting with the text wrapped as a clickable hyperlink.

    Creates a paragraph with:
    - A regular run containing the bullet character and tab
    - A hyperlink element containing a run with the prompt text

    The URL is built as: base_url + text (with spaces replaced by %20).

    Args:
        paragraph: The paragraph to format
        base_url: The base URL to prepend to the encoded text
        doc: The Document object (needed to create hyperlink relationships)
    """
    # Extract clean text (strip bullets, list markers, etc.)
    text = _extract_clean_text(paragraph.text)

    # Build URL: base_url + text with special characters encoded
    url = base_url + text.replace('&', '%26').replace(' ', '%20')

    # Clear all list formatting
    clear_all_list_formatting(paragraph)
    paragraph.style = 'Normal'
    clear_all_list_formatting(paragraph)
    paragraph.clear()

    # Add bullet character and tab as a regular run
    bullet_run = paragraph.add_run(BULLET_CHAR + '\t')
    bullet_run.font.name = FONT_NAME
    bullet_run.font.size = Pt(FONT_SIZE_PT)

    # Create hyperlink relationship
    r_id = doc.part.relate_to(
        url,
        'http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink',
        is_external=True
    )

    # Build hyperlink XML element
    hyperlink = OxmlElement('w:hyperlink')
    hyperlink.set(qn('r:id'), r_id)

    new_run = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')

    # Underline
    u = OxmlElement('w:u')
    u.set(qn('w:val'), 'single')
    rPr.append(u)

    # Blue color (standard hyperlink)
    c = OxmlElement('w:color')
    c.set(qn('w:val'), '0563C1')
    rPr.append(c)

    # Font name
    rFonts = OxmlElement('w:rFonts')
    rFonts.set(qn('w:ascii'), FONT_NAME)
    rFonts.set(qn('w:hAnsi'), FONT_NAME)
    rPr.append(rFonts)

    # Font size (half-points)
    sz = OxmlElement('w:sz')
    sz.set(qn('w:val'), str(FONT_SIZE_PT * 2))
    rPr.append(sz)
    szCs = OxmlElement('w:szCs')
    szCs.set(qn('w:val'), str(FONT_SIZE_PT * 2))
    rPr.append(szCs)

    new_run.append(rPr)

    t = OxmlElement('w:t')
    t.text = text
    new_run.append(t)

    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)

    # Final list formatting clear
    clear_all_list_formatting(paragraph)

    # Set indentation for bullet appearance
    paragraph.paragraph_format.left_indent = Inches(0.5)
    paragraph.paragraph_format.first_line_indent = Inches(-0.25)


def replace_cxo_in_document(doc, full_name: str):
    """
    Replace all occurrences of 'CXO' with the role's full name in every run.

    Args:
        doc: The Document object to modify in place
        full_name: The full role name to substitute (e.g., "Chief Information Officer")
    """
    for para in doc.paragraphs:
        for run in para.runs:
            if re.search(r'CXO', run.text, re.IGNORECASE):
                run.text = re.sub(r'CXO', full_name, run.text, flags=re.IGNORECASE)


def add_spacing_after(paragraph, points: int = 12):
    """
    Add spacing after a paragraph.

    Args:
        paragraph: The paragraph to add spacing to
        points: The amount of spacing in points (default 12)
    """
    paragraph.paragraph_format.space_after = Pt(points)


