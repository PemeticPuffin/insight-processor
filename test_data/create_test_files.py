#!/usr/bin/env python3
"""
Script to create complex test Word documents for manual testing of Insight Processor.

Run this script to generate test input files in the test_data/input folder.
"""

from pathlib import Path
from docx import Document
from docx.shared import Pt, Inches
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.enum.text import WD_ALIGN_PARAGRAPH


def add_hyperlink(paragraph, url, text):
    """Add a hyperlink to a paragraph."""
    part = paragraph.part
    r_id = part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True
    )

    hyperlink = OxmlElement('w:hyperlink')
    hyperlink.set(qn('r:id'), r_id)

    new_run = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')

    # Add underline
    u = OxmlElement('w:u')
    u.set(qn('w:val'), 'single')
    rPr.append(u)

    # Add blue color
    c = OxmlElement('w:color')
    c.set(qn('w:val'), '0000FF')
    rPr.append(c)

    new_run.append(rPr)

    t = OxmlElement('w:t')
    t.text = text
    new_run.append(t)

    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)

    return hyperlink


def create_insights_document(output_path: Path):
    """
    Create a complex insights document with multiple job roles.

    This document tests:
    - Multiple Heading 2 sections (job roles)
    - All header types that should be bolded
    - Bullet sections with intro lines to skip
    - Numbered and dashed lists to convert
    - Hyperlinks to preserve
    - Various formatting challenges
    """
    doc = Document()

    # Add a title/intro that should be ignored (before first H2)
    title = doc.add_paragraph("Weekly Insights Compilation")
    title.runs[0].bold = True
    title.runs[0].font.size = Pt(16)

    doc.add_paragraph("This document contains insights for multiple job roles. "
                      "Each section should be split into its own file.")
    doc.add_paragraph("")

    # ========== ROLE 1: Chief Audit Executive ==========
    doc.add_heading("Chief Audit Executive", level=2)

    doc.add_paragraph("Title of Insight: AI-Driven Audit Risk Assessment")

    doc.add_paragraph("Why This Matters")
    doc.add_paragraph("Key considerations for your organization:")  # Intro line - should NOT be bulleted
    doc.add_paragraph("1. Reduces manual effort in risk identification by up to 40%")
    doc.add_paragraph("2. Enables continuous monitoring instead of periodic reviews")
    doc.add_paragraph("3. Improves accuracy of risk scoring through pattern recognition")
    doc.add_paragraph("AI-powered audit tools are transforming how organizations approach internal controls.")

    doc.add_paragraph("Implications")
    doc.add_paragraph("Organizations that fail to adopt AI-driven audit capabilities risk falling behind "
                      "in an increasingly complex regulatory environment. Early adopters report 30% faster "
                      "audit cycles and improved stakeholder confidence.")

    doc.add_paragraph("Why Gartner")
    p = doc.add_paragraph("Gartner's research in AI for audit provides actionable frameworks. See our ")
    add_hyperlink(p, "https://www.gartner.com/en/audit-risk", "Audit & Risk Research Hub")
    p.add_run(" for detailed guidance.")

    doc.add_paragraph("Visual/Graphic")
    doc.add_paragraph("[Insert AI Audit Maturity Model graphic here]")

    doc.add_paragraph("Engage Clients and Prospects")
    doc.add_paragraph("Use this insight to open conversations about digital transformation in audit functions.")

    doc.add_paragraph("Probing Questions")
    doc.add_paragraph("Here are questions to drive the conversation:")  # Intro line
    doc.add_paragraph("- What percentage of your audit procedures are currently automated?")
    doc.add_paragraph("- How do you currently prioritize audit areas based on risk?")
    doc.add_paragraph("- What challenges have you faced with traditional sampling methods?")

    doc.add_paragraph("AskGartner Prompts to Recommend")
    doc.add_paragraph("The following prompts will help your clients:")  # Intro line
    doc.add_paragraph("a. What are the leading AI tools for internal audit?")
    doc.add_paragraph("b. How should we build an AI roadmap for our audit function?")
    doc.add_paragraph("c. What skills does my audit team need for AI adoption?")

    doc.add_paragraph("Insights to Engage")
    p = doc.add_paragraph("Related research: ")
    add_hyperlink(p, "https://www.gartner.com/document/4019234", "Top 10 Audit Technology Trends 2024")
    p.add_run(" and ")
    add_hyperlink(p, "https://www.gartner.com/document/4019567", "Building the AI-Ready Audit Team")

    # ========== ROLE 2: General Counsel ==========
    doc.add_heading("General Counsel", level=2)

    doc.add_paragraph("Title of Insight: Navigating AI Governance and Legal Liability")

    doc.add_paragraph("Why This Matters")
    doc.add_paragraph("Consider the following implications for legal departments:")  # Intro line
    doc.add_paragraph("Legal teams face unprecedented challenges as AI deployment accelerates across enterprises.")
    doc.add_paragraph("1) Regulatory frameworks are evolving rapidly across jurisdictions")
    doc.add_paragraph("2) Liability questions remain largely untested in courts")
    doc.add_paragraph("3) Contract language must adapt to address AI-specific risks")

    doc.add_paragraph("Implications")
    doc.add_paragraph("General Counsels who proactively establish AI governance frameworks position their "
                      "organizations to innovate safely. Those who wait may face costly remediation and "
                      "reputational damage when issues arise.")

    doc.add_paragraph("Why Gartner")
    p = doc.add_paragraph("Our ")
    add_hyperlink(p, "https://www.gartner.com/en/legal-compliance", "Legal & Compliance Research")
    p.add_run(" provides templates, frameworks, and peer benchmarks for AI governance.")

    doc.add_paragraph("Visual/Graphic")
    doc.add_paragraph("[Insert AI Legal Risk Framework diagram]")

    doc.add_paragraph("Engage Clients and Prospects")
    doc.add_paragraph("This insight resonates with legal leaders concerned about emerging technology risks.")

    doc.add_paragraph("Probing Questions")
    doc.add_paragraph("These questions help identify pain points:")  # Intro line
    doc.add_paragraph("(1) Does your organization have an AI ethics policy?")
    doc.add_paragraph("(2) How are you managing vendor AI liability in contracts?")
    doc.add_paragraph("(3) What process exists for legal review of AI deployments?")

    doc.add_paragraph("AskGartner Prompts to Recommend")
    doc.add_paragraph("Suggest these prompts:")  # Intro line
    doc.add_paragraph("How should legal structure AI governance oversight?")
    doc.add_paragraph("What contract clauses are essential for AI vendor agreements?")

    doc.add_paragraph("Insights to Engage")
    p = doc.add_paragraph("See also: ")
    add_hyperlink(p, "https://www.gartner.com/document/4020123", "AI Liability: What GCs Need to Know")

    # ========== ROLE 3: CPO - Procurement (GBS) ==========
    doc.add_heading("CPO - Procurement (GBS)", level=2)

    doc.add_paragraph("Title of Insight: Sustainable Procurement Through AI-Powered Supplier Intelligence")

    doc.add_paragraph("Why This Matters")
    doc.add_paragraph("Procurement leaders must address sustainability while maintaining cost efficiency:")  # Intro
    doc.add_paragraph("Supply chain sustainability is no longer optional - it's a board-level priority.")
    doc.add_paragraph("- 78% of enterprises now track Scope 3 emissions from suppliers")
    doc.add_paragraph("- Regulatory requirements (CSRD, SEC climate rules) demand supplier transparency")
    doc.add_paragraph("- AI can analyze thousands of suppliers for ESG risk in minutes")

    doc.add_paragraph("Implications")
    doc.add_paragraph("CPOs who leverage AI for supplier sustainability assessment gain competitive advantage "
                      "in attracting ESG-conscious customers and investors. Manual approaches cannot scale "
                      "to meet emerging disclosure requirements.")

    doc.add_paragraph("Why Gartner")
    doc.add_paragraph("Gartner's Sustainable Procurement Framework helps organizations balance cost, risk, "
                      "and sustainability objectives with practical implementation guidance.")

    doc.add_paragraph("Visual/Graphic")
    doc.add_paragraph("[Insert Supplier ESG Scoring Model visualization]")

    doc.add_paragraph("Engage Clients and Prospects")
    doc.add_paragraph("This is particularly relevant for organizations facing investor or regulatory "
                      "pressure on supply chain sustainability.")

    doc.add_paragraph("Probing Questions")
    doc.add_paragraph("Use these to understand their current state:")  # Intro
    doc.add_paragraph("How do you currently assess supplier sustainability performance?")
    doc.add_paragraph("What visibility do you have into Tier 2 and Tier 3 suppliers?")
    doc.add_paragraph("How are you preparing for upcoming ESG disclosure requirements?")

    doc.add_paragraph("AskGartner Prompts to Recommend")
    doc.add_paragraph("What are best practices for supplier sustainability scorecards?")
    doc.add_paragraph("How can AI improve our supplier risk monitoring?")

    doc.add_paragraph("Insights to Engage")
    p = doc.add_paragraph("Explore: ")
    add_hyperlink(p, "https://www.gartner.com/document/4021456", "Building Sustainable Supply Chains")

    # ========== ROLE 4: Tech CEO ==========
    doc.add_heading("Tech CEO", level=2)

    doc.add_paragraph("Title of Insight: Platform Strategy in the Age of AI Commoditization")

    doc.add_paragraph("Why This Matters")
    doc.add_paragraph("The rapid commoditization of AI capabilities creates both threats and opportunities "
                      "for technology companies. Leaders must rethink their competitive moats.")
    doc.add_paragraph("1. Foundation models are becoming commodity infrastructure")
    doc.add_paragraph("2. Differentiation is shifting to data, workflows, and ecosystem")
    doc.add_paragraph("3. Platform economics favor early movers with network effects")

    doc.add_paragraph("Implications")
    doc.add_paragraph("Tech CEOs who cling to proprietary AI models as differentiators risk disruption. "
                      "The winners will be those who build defensible data assets and ecosystem lock-in "
                      "while leveraging the best available AI infrastructure.")

    doc.add_paragraph("Why Gartner")
    p = doc.add_paragraph("Gartner's ")
    add_hyperlink(p, "https://www.gartner.com/en/information-technology", "Technology CEO Research")
    p.add_run(" provides strategic frameworks for navigating AI market dynamics.")

    doc.add_paragraph("Visual/Graphic")
    doc.add_paragraph("[Insert AI Platform Strategy Matrix]")

    doc.add_paragraph("Engage Clients and Prospects")
    doc.add_paragraph("This insight sparks strategic conversations about long-term competitive positioning.")

    doc.add_paragraph("Probing Questions")
    doc.add_paragraph("Below are questions to challenge strategic thinking:")  # Intro
    doc.add_paragraph("What is your sustainable competitive advantage as AI commoditizes?")
    doc.add_paragraph("How are you building data and ecosystem moats?")
    doc.add_paragraph("What is your build vs. buy strategy for AI capabilities?")

    doc.add_paragraph("AskGartner Prompts to Recommend")
    doc.add_paragraph("How should tech companies position against AI hyperscalers?")
    doc.add_paragraph("What platform strategies create durable AI advantages?")

    doc.add_paragraph("Insights to Engage")
    p = doc.add_paragraph("Related: ")
    add_hyperlink(p, "https://www.gartner.com/document/4022789", "The Future of Tech Platform Competition")

    # ========== ROLE 5: R&D ==========
    doc.add_heading("R&D", level=2)

    doc.add_paragraph("Title of Insight: Accelerating Drug Discovery with Generative AI")

    doc.add_paragraph("Why This Matters")
    doc.add_paragraph("Here is why R&D leaders should prioritize generative AI:")  # Intro
    doc.add_paragraph("Generative AI is revolutionizing pharmaceutical R&D timelines and success rates.")
    doc.add_paragraph("- Molecule generation: 10x faster candidate identification")
    doc.add_paragraph("- Clinical trial design: AI-optimized protocols reduce patient burden")
    doc.add_paragraph("- Literature synthesis: Automated review of thousands of papers in hours")

    doc.add_paragraph("Implications")
    doc.add_paragraph("R&D organizations that integrate generative AI into their workflows will significantly "
                      "outpace competitors in time-to-market. The technology is mature enough for production "
                      "use in specific applications.")

    doc.add_paragraph("Why Gartner")
    doc.add_paragraph("Our Life Sciences research team tracks AI adoption patterns and provides "
                      "vendor landscape analysis for R&D technology investments.")

    doc.add_paragraph("Visual/Graphic")
    doc.add_paragraph("[Insert AI in Drug Discovery Pipeline graphic]")

    doc.add_paragraph("Engage Clients and Prospects")
    doc.add_paragraph("This insight is highly relevant for pharma and biotech R&D executives exploring AI.")

    doc.add_paragraph("Probing Questions")
    doc.add_paragraph("Questions to assess AI readiness:")  # Intro
    doc.add_paragraph("What AI/ML capabilities exist in your current R&D tech stack?")
    doc.add_paragraph("How are you evaluating generative AI vendors for drug discovery?")
    doc.add_paragraph("What data infrastructure gaps limit your AI potential?")

    doc.add_paragraph("AskGartner Prompts to Recommend")
    doc.add_paragraph("Which generative AI tools are most mature for pharma R&D?")
    doc.add_paragraph("How should we structure AI governance for regulated research?")

    doc.add_paragraph("Insights to Engage")
    p = doc.add_paragraph("See: ")
    add_hyperlink(p, "https://www.gartner.com/document/4023456", "Generative AI in Life Sciences: A Practical Guide")

    # Save the document
    doc.save(str(output_path))
    print(f"Created: {output_path}")


def create_email_templates_document(output_path: Path):
    """
    Create a complex email templates document.

    This document tests:
    - Multiple job roles with Heading 2
    - BD, AE, EP email types with Heading 3
    - Various content and formatting
    - Role abbreviation mapping
    """
    doc = Document()

    # Title
    title = doc.add_paragraph("Email Templates - Week of January 27, 2025")
    title.runs[0].bold = True
    title.runs[0].font.size = Pt(16)

    doc.add_paragraph("Use these templates for outreach. Customize based on client context.")
    doc.add_paragraph("")

    # ========== Chief Audit Executive ==========
    doc.add_heading("Chief Audit Executive", level=2)

    doc.add_heading("SALES BUSINESS DEVELOPER PROSPECT EMAIL", level=3)
    doc.add_paragraph("Subject: Transform Your Audit Function with AI-Powered Risk Assessment")
    doc.add_paragraph("")
    doc.add_paragraph("Dear [Prospect Name],")
    doc.add_paragraph("")
    doc.add_paragraph("I hope this message finds you well. As organizations face increasingly complex risk "
                      "landscapes, I wanted to share how leading audit functions are leveraging AI to "
                      "enhance their capabilities.")
    doc.add_paragraph("")
    doc.add_paragraph("Gartner research shows that organizations using AI-driven audit tools experience:")
    doc.add_paragraph("- 40% reduction in manual risk identification effort")
    doc.add_paragraph("- 30% faster audit cycle times")
    doc.add_paragraph("- Improved accuracy in risk scoring and prioritization")
    doc.add_paragraph("")
    doc.add_paragraph("Would you be open to a 15-minute conversation about how these trends might apply "
                      "to your organization?")
    doc.add_paragraph("")
    doc.add_paragraph("Best regards,")
    doc.add_paragraph("[Your Name]")

    doc.add_heading("SALES ACCOUNT EXECUTIVE CLIENT EMAIL", level=3)
    doc.add_paragraph("Subject: Following Up on AI Audit Discussion")
    doc.add_paragraph("")
    doc.add_paragraph("Dear [Client Name],")
    doc.add_paragraph("")
    doc.add_paragraph("Thank you for our recent conversation about your audit transformation initiatives. "
                      "Based on our discussion, I wanted to share some additional resources that may be helpful:")
    doc.add_paragraph("")
    doc.add_paragraph("1. Top 10 Audit Technology Trends 2024 - Key insights on AI adoption")
    doc.add_paragraph("2. Building the AI-Ready Audit Team - Skills and capability framework")
    doc.add_paragraph("3. AI Audit Maturity Model - Assess your current state")
    doc.add_paragraph("")
    doc.add_paragraph("I'd also like to introduce you to our audit research team who can provide "
                      "deeper guidance on your specific challenges.")
    doc.add_paragraph("")
    doc.add_paragraph("Would next Tuesday or Wednesday work for a follow-up call?")
    doc.add_paragraph("")
    doc.add_paragraph("Best regards,")
    doc.add_paragraph("[Your Name]")

    doc.add_heading("SERVICE EXECUTIVE PARTNER", level=3)
    doc.add_paragraph("Subject: Strategic Partnership for Audit Excellence")
    doc.add_paragraph("")
    doc.add_paragraph("Dear [Client Name],")
    doc.add_paragraph("")
    doc.add_paragraph("As your Executive Partner, I wanted to ensure you're getting maximum value "
                      "from your Gartner relationship in the audit and risk domain.")
    doc.add_paragraph("")
    doc.add_paragraph("Based on your strategic priorities, I recommend we schedule a session with our "
                      "Chief Audit Executive research team to discuss:")
    doc.add_paragraph("")
    doc.add_paragraph("- Your AI adoption roadmap for audit")
    doc.add_paragraph("- Peer benchmarks from similar organizations")
    doc.add_paragraph("- Actionable next steps aligned with your timeline")
    doc.add_paragraph("")
    doc.add_paragraph("I'll coordinate with your team to find a suitable time. Please let me know "
                      "any specific topics you'd like to prioritize.")
    doc.add_paragraph("")
    doc.add_paragraph("Warm regards,")
    doc.add_paragraph("[Your Name]")

    # ========== General Counsel ==========
    doc.add_heading("General Counsel", level=2)

    doc.add_heading("SALES BUSINESS DEVELOPER PROSPECT EMAIL", level=3)
    doc.add_paragraph("Subject: AI Governance - A Critical Priority for Legal Leaders")
    doc.add_paragraph("")
    doc.add_paragraph("Dear [Prospect Name],")
    doc.add_paragraph("")
    doc.add_paragraph("With AI deployment accelerating across enterprises, General Counsels face "
                      "unprecedented governance and liability challenges.")
    doc.add_paragraph("")
    doc.add_paragraph("Gartner research indicates that organizations with proactive AI governance "
                      "frameworks experience 60% fewer compliance incidents and are better positioned "
                      "for upcoming regulations like the EU AI Act.")
    doc.add_paragraph("")
    doc.add_paragraph("I'd welcome the opportunity to share insights on how peer organizations are "
                      "approaching AI governance in their legal functions.")
    doc.add_paragraph("")
    doc.add_paragraph("Do you have 15 minutes this week for a brief conversation?")
    doc.add_paragraph("")
    doc.add_paragraph("Best regards,")
    doc.add_paragraph("[Your Name]")

    doc.add_heading("SALES ACCOUNT EXECUTIVE CLIENT EMAIL", level=3)
    doc.add_paragraph("Subject: Resources for Your AI Governance Initiative")
    doc.add_paragraph("")
    doc.add_paragraph("Dear [Client Name],")
    doc.add_paragraph("")
    doc.add_paragraph("Following our discussion about AI governance, I wanted to share several "
                      "resources that address your key concerns:")
    doc.add_paragraph("")
    doc.add_paragraph("- AI Liability: What GCs Need to Know")
    doc.add_paragraph("- Contract Clauses for AI Vendor Agreements")
    doc.add_paragraph("- Building an AI Ethics Framework")
    doc.add_paragraph("")
    doc.add_paragraph("Our legal research analysts are available for a deep-dive session on any "
                      "of these topics.")
    doc.add_paragraph("")
    doc.add_paragraph("Best regards,")
    doc.add_paragraph("[Your Name]")

    doc.add_heading("SERVICE EXECUTIVE PARTNER", level=3)
    doc.add_paragraph("Subject: Maximizing Your Legal Research Engagement")
    doc.add_paragraph("")
    doc.add_paragraph("Dear [Client Name],")
    doc.add_paragraph("")
    doc.add_paragraph("I wanted to check in on your experience with our Legal & Compliance research "
                      "and identify opportunities to drive additional value.")
    doc.add_paragraph("")
    doc.add_paragraph("Given the rapidly evolving AI regulatory landscape, I recommend we arrange "
                      "a strategic session to align your Gartner engagement with your 2025 priorities.")
    doc.add_paragraph("")
    doc.add_paragraph("Looking forward to connecting.")
    doc.add_paragraph("")
    doc.add_paragraph("Warm regards,")
    doc.add_paragraph("[Your Name]")

    # ========== CPO - Procurement (GBS) ==========
    doc.add_heading("CPO - Procurement (GBS)", level=2)

    doc.add_heading("SALES BUSINESS DEVELOPER PROSPECT EMAIL", level=3)
    doc.add_paragraph("Subject: Sustainable Procurement - Meeting ESG Mandates")
    doc.add_paragraph("")
    doc.add_paragraph("Dear [Prospect Name],")
    doc.add_paragraph("")
    doc.add_paragraph("With increasing pressure from regulators and investors on supply chain "
                      "sustainability, procurement leaders are seeking new approaches to ESG compliance.")
    doc.add_paragraph("")
    doc.add_paragraph("Gartner research shows that AI-powered supplier intelligence can reduce "
                      "ESG assessment time by 80% while improving accuracy.")
    doc.add_paragraph("")
    doc.add_paragraph("Would you be interested in learning how leading CPOs are tackling "
                      "sustainable procurement challenges?")
    doc.add_paragraph("")
    doc.add_paragraph("Best regards,")
    doc.add_paragraph("[Your Name]")

    doc.add_heading("SALES ACCOUNT EXECUTIVE CLIENT EMAIL", level=3)
    doc.add_paragraph("Subject: Advancing Your Sustainable Procurement Strategy")
    doc.add_paragraph("")
    doc.add_paragraph("Dear [Client Name],")
    doc.add_paragraph("")
    doc.add_paragraph("Thank you for sharing your sustainable procurement priorities. Based on "
                      "our conversation, here are resources I recommend:")
    doc.add_paragraph("")
    doc.add_paragraph("- Building Sustainable Supply Chains")
    doc.add_paragraph("- Supplier ESG Scorecard Best Practices")
    doc.add_paragraph("- AI for Supplier Risk Monitoring")
    doc.add_paragraph("")
    doc.add_paragraph("Let me know if you'd like to schedule an analyst inquiry on any of these topics.")
    doc.add_paragraph("")
    doc.add_paragraph("Best regards,")
    doc.add_paragraph("[Your Name]")

    doc.add_heading("SERVICE EXECUTIVE PARTNER", level=3)
    doc.add_paragraph("Subject: Quarterly Business Review - Procurement")
    doc.add_paragraph("")
    doc.add_paragraph("Dear [Client Name],")
    doc.add_paragraph("")
    doc.add_paragraph("As we approach the end of Q1, I'd like to schedule our quarterly review "
                      "to assess progress on your procurement priorities.")
    doc.add_paragraph("")
    doc.add_paragraph("Suggested agenda:")
    doc.add_paragraph("- Progress on sustainable procurement initiatives")
    doc.add_paragraph("- Upcoming research and events")
    doc.add_paragraph("- 2025 strategic priorities alignment")
    doc.add_paragraph("")
    doc.add_paragraph("Please let me know your availability for a 45-minute session.")
    doc.add_paragraph("")
    doc.add_paragraph("Warm regards,")
    doc.add_paragraph("[Your Name]")

    # ========== Tech CEO ==========
    doc.add_heading("Tech CEO", level=2)

    doc.add_heading("SALES BUSINESS DEVELOPER PROSPECT EMAIL", level=3)
    doc.add_paragraph("Subject: Platform Strategy in the AI Era")
    doc.add_paragraph("")
    doc.add_paragraph("Dear [Prospect Name],")
    doc.add_paragraph("")
    doc.add_paragraph("The rapid commoditization of AI capabilities is reshaping competitive dynamics "
                      "for technology companies. Gartner research identifies key strategies for "
                      "maintaining differentiation.")
    doc.add_paragraph("")
    doc.add_paragraph("I'd welcome the opportunity to discuss how leading tech CEOs are "
                      "rethinking their platform strategies.")
    doc.add_paragraph("")
    doc.add_paragraph("Would you have 20 minutes for a conversation?")
    doc.add_paragraph("")
    doc.add_paragraph("Best regards,")
    doc.add_paragraph("[Your Name]")

    doc.add_heading("SALES ACCOUNT EXECUTIVE CLIENT EMAIL", level=3)
    doc.add_paragraph("Subject: Strategic Resources for Platform Competition")
    doc.add_paragraph("")
    doc.add_paragraph("Dear [Client Name],")
    doc.add_paragraph("")
    doc.add_paragraph("Following our platform strategy discussion, I wanted to share relevant research:")
    doc.add_paragraph("")
    doc.add_paragraph("- The Future of Tech Platform Competition")
    doc.add_paragraph("- Building Durable AI Advantages")
    doc.add_paragraph("- Data Moat Strategies for Tech Companies")
    doc.add_paragraph("")
    doc.add_paragraph("Let me know if you'd like to arrange an analyst session.")
    doc.add_paragraph("")
    doc.add_paragraph("Best regards,")
    doc.add_paragraph("[Your Name]")

    doc.add_heading("SERVICE EXECUTIVE PARTNER", level=3)
    doc.add_paragraph("Subject: Executive Briefing Opportunity")
    doc.add_paragraph("")
    doc.add_paragraph("Dear [Client Name],")
    doc.add_paragraph("")
    doc.add_paragraph("I wanted to extend an invitation for an executive briefing with our "
                      "Tech CEO research leadership.")
    doc.add_paragraph("")
    doc.add_paragraph("This session would cover:")
    doc.add_paragraph("- AI market dynamics and competitive implications")
    doc.add_paragraph("- Platform strategy recommendations")
    doc.add_paragraph("- Peer insights from similar tech companies")
    doc.add_paragraph("")
    doc.add_paragraph("Let me know if this would be valuable for your leadership team.")
    doc.add_paragraph("")
    doc.add_paragraph("Warm regards,")
    doc.add_paragraph("[Your Name]")

    # ========== R&D ==========
    doc.add_heading("R&D", level=2)

    doc.add_heading("SALES BUSINESS DEVELOPER PROSPECT EMAIL", level=3)
    doc.add_paragraph("Subject: Generative AI in Drug Discovery - Practical Applications")
    doc.add_paragraph("")
    doc.add_paragraph("Dear [Prospect Name],")
    doc.add_paragraph("")
    doc.add_paragraph("Generative AI is transforming pharmaceutical R&D timelines. Leading "
                      "organizations report 10x faster candidate identification and significantly "
                      "improved success rates.")
    doc.add_paragraph("")
    doc.add_paragraph("Would you be interested in learning how peers are implementing "
                      "generative AI in their R&D workflows?")
    doc.add_paragraph("")
    doc.add_paragraph("Best regards,")
    doc.add_paragraph("[Your Name]")

    doc.add_heading("SALES ACCOUNT EXECUTIVE CLIENT EMAIL", level=3)
    doc.add_paragraph("Subject: R&D AI Resources")
    doc.add_paragraph("")
    doc.add_paragraph("Dear [Client Name],")
    doc.add_paragraph("")
    doc.add_paragraph("Based on our conversation about AI in drug discovery, here are key resources:")
    doc.add_paragraph("")
    doc.add_paragraph("- Generative AI in Life Sciences: A Practical Guide")
    doc.add_paragraph("- AI Vendor Landscape for Pharma R&D")
    doc.add_paragraph("- Building AI Governance for Regulated Research")
    doc.add_paragraph("")
    doc.add_paragraph("Let me know your thoughts.")
    doc.add_paragraph("")
    doc.add_paragraph("Best regards,")
    doc.add_paragraph("[Your Name]")

    doc.add_heading("SERVICE EXECUTIVE PARTNER", level=3)
    doc.add_paragraph("Subject: R&D Innovation Partnership")
    doc.add_paragraph("")
    doc.add_paragraph("Dear [Client Name],")
    doc.add_paragraph("")
    doc.add_paragraph("I'd like to explore how we can better support your R&D innovation agenda "
                      "through your Gartner partnership.")
    doc.add_paragraph("")
    doc.add_paragraph("Suggested topics for our next strategic session:")
    doc.add_paragraph("- AI/ML technology roadmap")
    doc.add_paragraph("- Vendor evaluation support")
    doc.add_paragraph("- Peer networking opportunities")
    doc.add_paragraph("")
    doc.add_paragraph("Looking forward to connecting.")
    doc.add_paragraph("")
    doc.add_paragraph("Warm regards,")
    doc.add_paragraph("[Your Name]")

    # Save the document
    doc.save(str(output_path))
    print(f"Created: {output_path}")


def create_simple_insights_document(output_path: Path):
    """
    Create a simpler insights document for quick testing.

    Just 2 roles with basic content.
    """
    doc = Document()

    # ========== ROLE 1: Customer Service ==========
    doc.add_heading("Customer Service", level=2)

    doc.add_paragraph("Title of Insight: AI-Powered Customer Support Transformation")

    doc.add_paragraph("Why This Matters")
    doc.add_paragraph("Customer expectations for support are rising while budgets remain constrained.")
    doc.add_paragraph("- AI chatbots now handle 45% of routine inquiries")
    doc.add_paragraph("- Response times reduced by 60% with AI triage")
    doc.add_paragraph("- Customer satisfaction scores improve with faster resolution")

    doc.add_paragraph("Implications")
    doc.add_paragraph("Organizations that delay AI adoption in customer service risk losing customers "
                      "to competitors offering superior digital experiences.")

    doc.add_paragraph("Probing Questions")
    doc.add_paragraph("What is your current AI adoption level in customer service?")
    doc.add_paragraph("How do you measure customer effort score?")

    doc.add_paragraph("AskGartner Prompts to Recommend")
    doc.add_paragraph("What are best practices for AI-human handoff in support?")

    # ========== ROLE 2: Unknown Role (tests fallback) ==========
    doc.add_heading("VP of Digital Innovation", level=2)

    doc.add_paragraph("Title of Insight: Building a Digital Innovation Lab")

    doc.add_paragraph("Why This Matters")
    doc.add_paragraph("Innovation labs are evolving from experiments to enterprise capabilities.")
    doc.add_paragraph("1. Labs must demonstrate business value within 12-18 months")
    doc.add_paragraph("2. Successful labs integrate with core business operations")
    doc.add_paragraph("3. Talent strategy is critical for sustainability")

    doc.add_paragraph("Implications")
    doc.add_paragraph("Innovation leaders who align lab initiatives with strategic priorities "
                      "achieve higher success rates and executive support.")

    doc.add_paragraph("Probing Questions")
    doc.add_paragraph("How is your innovation lab connected to business units?")
    doc.add_paragraph("What metrics define success for your lab?")

    doc.add_paragraph("AskGartner Prompts to Recommend")
    doc.add_paragraph("How should we structure innovation lab governance?")

    doc.save(str(output_path))
    print(f"Created: {output_path}")


def main():
    """Generate all test files."""
    script_dir = Path(__file__).parent
    input_dir = script_dir / "input"
    input_dir.mkdir(exist_ok=True)

    print("Creating test Word documents for manual testing...\n")

    # Create main insights document (complex, 5 roles)
    create_insights_document(input_dir / "Weekly_Insights_Full.docx")

    # Create email templates document
    create_email_templates_document(input_dir / "Email_Templates_Full.docx")

    # Create simpler insights document (2 roles)
    create_simple_insights_document(input_dir / "Weekly_Insights_Simple.docx")

    print("\n" + "=" * 60)
    print("TEST FILES CREATED!")
    print("=" * 60)
    print(f"\nInput files are in: {input_dir}")
    print(f"Output folder ready: {script_dir / 'output'}")
    print("\nTo test the application:")
    print("  1. Run: python insight_processor.py")
    print("  2. Choose options from the menu")
    print(f"  3. Point to input files in: {input_dir}")
    print(f"  4. Point to output folder: {script_dir / 'output'}")
    print("\nTest scenarios:")
    print("  - Step 1: Split Weekly_Insights_Full.docx (5 role subfolders created)")
    print("  - Step 2: Clean/rename files in output folder")
    print("  - Step 3: Format all insight files")
    print("  - Step 4: Split Email_Templates_Full.docx into role folders")
    print("  - Option A: Run all steps in sequence")


if __name__ == "__main__":
    main()
