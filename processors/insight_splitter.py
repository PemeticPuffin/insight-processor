"""
Step 1: Split Insights by Heading 2

Splits a Word document into separate files based on Heading 2 sections.
Creates files in job-role subfolders, creating subfolders if they don't exist.
Preserves all hyperlinks using relationship mapping.
"""

from pathlib import Path
from copy import deepcopy
from docx import Document
from typing import Dict, List, Tuple, Any

import sys
sys.path.insert(0, str(Path(__file__).parent.parent))

from utils.paragraph_utils import copy_element_relationships
from utils.folder_utils import ensure_subfolder_exists


def is_heading_2(paragraph) -> bool:
    """Check if a paragraph is styled as Heading 2."""
    return paragraph.style.name == 'Heading 2'


def extract_sections(doc: Document) -> List[Dict[str, Any]]:
    """
    Extract all Heading 2 sections from a document.

    Args:
        doc: The Document object to extract from

    Returns:
        List of dicts with 'heading' and 'content' keys
    """
    current_section = None
    current_content = []
    sections = []

    for paragraph in doc.paragraphs:
        if is_heading_2(paragraph):
            # Save previous section if it exists
            if current_section is not None:
                sections.append({
                    'heading': current_section,
                    'content': current_content
                })

            # Start new section
            current_section = paragraph.text.strip()
            current_content = [paragraph]
        else:
            # Add content to current section
            if current_section is not None:
                current_content.append(paragraph)

    # Save the last section
    if current_section is not None:
        sections.append({
            'heading': current_section,
            'content': current_content
        })

    return sections


def create_section_document(section: Dict[str, Any], source_doc: Document) -> Document:
    """
    Create a new document from a section's content.

    Args:
        section: Dict with 'heading' and 'content' keys
        source_doc: The source Document for relationship copying

    Returns:
        New Document containing the section content
    """
    new_doc = Document()

    for para in section['content']:
        # Deep copy the paragraph element to preserve hyperlinks and all formatting
        new_para_element = deepcopy(para._element)

        # Copy all relationships (hyperlinks, images, etc.) before adding
        copy_element_relationships(new_para_element, source_doc.part, new_doc.part)

        new_doc._body._body.append(new_para_element)

    return new_doc


def split_insights(
    input_file: Path,
    output_directory: Path,
    create_subfolders: bool = True,
    progress_callback=None
) -> Tuple[int, int, List[str]]:
    """
    Split a Word document by Heading 2 sections.

    Args:
        input_file: Path to the input Word document
        output_directory: Directory to save output files (with subfolders)
        create_subfolders: If True, create subfolders that don't exist
        progress_callback: Optional callback function(message: str)

    Returns:
        Tuple of (files_created, sections_found, list of created file paths)
    """
    def log(msg):
        if progress_callback:
            progress_callback(msg)
        else:
            print(msg)

    # Load the document
    doc = Document(input_file)

    # Extract sections
    sections = extract_sections(doc)

    if not sections:
        log("No Heading 2 sections found in the document.")
        return 0, 0, []

    log(f"Found {len(sections)} Heading 2 sections")

    # Create documents for each section
    files_created = 0
    created_files = []

    for section in sections:
        heading = section['heading']

        # Handle subfolder
        if create_subfolders:
            subfolder_path, was_created = ensure_subfolder_exists(output_directory, heading)
            action = "Created" if was_created else "Using existing"
            log(f"  {action} subfolder: {subfolder_path.name}")
        else:
            subfolder_path = output_directory / heading
            if not subfolder_path.exists():
                log(f"  Warning: Subfolder '{heading}' not found. Skipping this section.")
                continue

        # Create new document with section content
        new_doc = create_section_document(section, doc)

        # Use the heading text as the filename
        output_filename = f"{heading}.docx"
        output_filepath = subfolder_path / output_filename

        # Save the new document
        new_doc.save(str(output_filepath))

        log(f"  Created: {output_filepath.name}")
        files_created += 1
        created_files.append(str(output_filepath))

    return files_created, len(sections), created_files


def get_heading_list(input_file: Path) -> List[str]:
    """
    Get a list of all Heading 2 texts from a document.

    Args:
        input_file: Path to the Word document

    Returns:
        List of Heading 2 text strings
    """
    doc = Document(input_file)
    return [p.text.strip() for p in doc.paragraphs if is_heading_2(p)]
