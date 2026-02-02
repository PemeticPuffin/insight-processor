"""
Step 4: Split Email Templates

Splits email templates from a Word document into separate files.
Creates BD_AE_emails.docx and EP_email.docx per job role.
Applies Calibri 11pt formatting to all output files.
"""

from pathlib import Path
from copy import deepcopy
from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn
from typing import Dict, List, Tuple, Any, Optional

import sys
sys.path.insert(0, str(Path(__file__).parent.parent))

from config import FONT_NAME, FONT_SIZE_PT
from utils.date_utils import get_two_mondays_from_now, format_date_for_filename
from utils.text_utils import get_role_abbreviation
from utils.folder_utils import find_matching_subfolder

# XML namespace for relationships
HYPERLINK_TAG = qn('w:hyperlink')


def apply_font_formatting(doc: Document):
    """
    Apply Calibri 11pt font to all paragraphs in a document.

    Args:
        doc: The Document to format
    """
    for para in doc.paragraphs:
        for run in para.runs:
            run.font.name = FONT_NAME
            run.font.size = Pt(FONT_SIZE_PT)


def copy_paragraph_with_formatting(source_para, target_doc, source_doc=None) -> Any:
    """
    Copy a paragraph with all its formatting to target document.

    Properly handles hyperlinks by re-creating relationships in the target document.

    Args:
        source_para: The source paragraph
        target_doc: The target Document object
        source_doc: The source Document object (needed for hyperlink relationships)

    Returns:
        The new paragraph in the target document
    """
    # Create new paragraph
    new_para = target_doc.add_paragraph()

    # Copy paragraph style
    try:
        new_para.style = source_para.style
    except Exception:
        pass

    # Copy paragraph formatting
    pf = source_para.paragraph_format
    npf = new_para.paragraph_format

    if pf.alignment is not None:
        npf.alignment = pf.alignment
    if pf.left_indent is not None:
        npf.left_indent = pf.left_indent
    if pf.right_indent is not None:
        npf.right_indent = pf.right_indent
    if pf.first_line_indent is not None:
        npf.first_line_indent = pf.first_line_indent
    if pf.space_before is not None:
        npf.space_before = pf.space_before
    if pf.space_after is not None:
        npf.space_after = pf.space_after

    # Clear and copy content
    new_para._element.clear_content()

    for child in source_para._element:
        child_copy = deepcopy(child)

        # Handle hyperlinks - need to re-create relationship in target document
        if child_copy.tag == HYPERLINK_TAG and source_doc is not None:
            # Get the relationship ID from the hyperlink
            r_id = child_copy.get(qn('r:id'))
            if r_id:
                # Look up the URL in the source document's relationships
                try:
                    source_part = source_doc.part
                    rel = source_part.rels.get(r_id)
                    if rel and rel.target_ref:
                        # Add new relationship in target document
                        new_rel = target_doc.part.relate_to(
                            rel.target_ref,
                            'http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink',
                            is_external=True
                        )
                        # Update the hyperlink element with the new relationship ID
                        child_copy.set(qn('r:id'), new_rel)
                except Exception:
                    # If relationship lookup fails, remove the r:id to avoid corruption
                    if child_copy.get(qn('r:id')):
                        del child_copy.attrib[qn('r:id')]

        new_para._element.append(child_copy)

    return new_para


def parse_email_document(doc_path: Path) -> Tuple[Dict[str, Dict[str, Any]], Document]:
    """
    Parse the Word document and extract job sections with email templates.

    Args:
        doc_path: Path to the email templates document

    Returns:
        Tuple of (Dict mapping job titles to their email templates, source Document)
    """
    doc = Document(doc_path)
    sections = {}
    current_job = None
    current_heading3 = None

    for para in doc.paragraphs:
        # Check if it's a Heading 2 (job title)
        if para.style.name == 'Heading 2':
            current_job = para.text.strip()
            sections[current_job] = {
                'BD': {'heading': None, 'content': []},
                'AE': {'heading': None, 'content': []},
                'EP': {'heading': None, 'content': []}
            }
            current_heading3 = None

        # Check if it's a Heading 3 (email template type)
        elif para.style.name == 'Heading 3' and current_job:
            heading_text = para.text.strip().upper()

            # Check for BD (Sales Business Developer / Prospect Email)
            if 'SALES BUSINESS DEVELOPER' in heading_text or 'PROSPECT EMAIL' in heading_text:
                current_heading3 = 'BD'
                sections[current_job]['BD']['heading'] = para
            # Check for AE (Sales Account Executive)
            elif 'SALES ACCOUNT EXECUTIVE' in heading_text:
                current_heading3 = 'AE'
                sections[current_job]['AE']['heading'] = para
            # Check for EP (Service Executive Partner) - must check BEFORE generic CLIENT EMAIL
            elif 'SERVICE EXECUTIVE PARTNER' in heading_text:
                current_heading3 = 'EP'
                sections[current_job]['EP']['heading'] = para
            # Check for EP (Service Client Email)
            elif 'SERVICE' in heading_text and 'CLIENT' in heading_text:
                current_heading3 = 'EP'
                sections[current_job]['EP']['heading'] = para
            # Check for AE (Client Email without Sales or Service)
            elif 'CLIENT EMAIL' in heading_text and 'SALES' not in heading_text and 'SERVICE' not in heading_text:
                current_heading3 = 'AE'
                sections[current_job]['AE']['heading'] = para

        # Add content to current section
        elif current_heading3 and current_job:
            sections[current_job][current_heading3]['content'].append(para)

    return sections, doc


def create_email_document(
    heading_paras: List[Any],
    content_paras_list: List[List[Any]],
    source_doc: Document = None
) -> Document:
    """
    Create a new Word document with headings and content.

    Args:
        heading_paras: List of heading paragraphs
        content_paras_list: List of lists of content paragraphs
        source_doc: Source Document object (needed for hyperlink relationships)

    Returns:
        New Document with the combined content
    """
    doc = Document()

    for heading_para, content_paras in zip(heading_paras, content_paras_list):
        # Add heading
        if heading_para:
            copy_paragraph_with_formatting(heading_para, doc, source_doc)

        # Add content
        for para in content_paras:
            copy_paragraph_with_formatting(para, doc, source_doc)

    # Apply Calibri 11pt formatting
    apply_font_formatting(doc)

    return doc


def split_emails(
    source_doc_path: Path,
    target_folder: Path,
    progress_callback=None
) -> Tuple[int, int, List[str]]:
    """
    Split email templates into separate files per job role.

    Args:
        source_doc_path: Path to the email templates document
        target_folder: Directory containing job-role subfolders
        progress_callback: Optional callback function(message: str)

    Returns:
        Tuple of (jobs_processed, jobs_skipped, list of created file paths)
    """
    def log(msg):
        if progress_callback:
            progress_callback(msg)
        else:
            print(msg)

    log(f"Reading source document: {source_doc_path.name}")
    sections, source_doc = parse_email_document(source_doc_path)

    # Calculate date prefix
    target_date = get_two_mondays_from_now()
    date_prefix = format_date_for_filename(target_date)
    log(f"Using date prefix: {date_prefix}")

    if not target_folder.exists():
        log(f"Error: Target folder does not exist: {target_folder}")
        return 0, 0, []

    processed_count = 0
    skipped_count = 0
    created_files = []

    for job_title, templates in sections.items():
        role = get_role_abbreviation(job_title)

        # Find matching subfolder (case-insensitive)
        subfolder = find_matching_subfolder(target_folder, job_title)

        if not subfolder:
            log(f"Skipping '{job_title}': subfolder not found")
            skipped_count += 1
            continue

        log(f"Processing: {job_title} -> {role}")

        # Create BD_AE_emails document
        if templates['BD']['heading'] or templates['AE']['heading']:
            bd_ae_doc = create_email_document(
                [templates['BD']['heading'], templates['AE']['heading']],
                [templates['BD']['content'], templates['AE']['content']],
                source_doc
            )
            bd_ae_filename = f"{date_prefix}_{role}_BD_AE_emails.docx"
            bd_ae_path = subfolder / bd_ae_filename

            bd_ae_doc.save(str(bd_ae_path))
            log(f"  Created: {bd_ae_filename}")
            created_files.append(str(bd_ae_path))

        # Create EP_email document
        if templates['EP']['heading']:
            ep_doc = create_email_document(
                [templates['EP']['heading']],
                [templates['EP']['content']],
                source_doc
            )
            ep_filename = f"{date_prefix}_{role}_EP_email.docx"
            ep_path = subfolder / ep_filename

            ep_doc.save(str(ep_path))
            log(f"  Created: {ep_filename}")
            created_files.append(str(ep_path))

        processed_count += 1

    return processed_count, skipped_count, created_files
