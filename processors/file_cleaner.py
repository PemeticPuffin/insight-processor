"""
Step 2: Clean and Rename Files

Processes Word documents in subfolders:
1. Delete the first line of text
2. Extract the 3rd line for the new filename
3. Rename using: [date 2 Mondays from now]_[subfolder]_[3rd line text].docx
"""

from pathlib import Path
from docx import Document
from typing import Tuple, Optional

import sys
sys.path.insert(0, str(Path(__file__).parent.parent))

from utils.date_utils import get_two_mondays_from_now, format_date_for_filename
from utils.text_utils import sanitize_for_filename


def find_matching_word_file(subfolder_path: Path, subfolder_name: str) -> Optional[Path]:
    """
    Find a Word document that matches the subfolder name.

    Args:
        subfolder_path: Path to the subfolder
        subfolder_name: Name of the subfolder to match

    Returns:
        Path to the matching file, or None if not found
    """
    # Look for .docx files
    for ext in ['.docx']:
        potential_file = subfolder_path / f"{subfolder_name}{ext}"
        if potential_file.exists():
            return potential_file
    return None


def process_single_document(
    doc_path: Path,
    subfolder_name: str,
    target_date_str: str
) -> Tuple[bool, str, Optional[str]]:
    """
    Process a single Word document: delete first line, extract 3rd line, rename.

    Args:
        doc_path: Path to the document
        subfolder_name: Name of the containing subfolder
        target_date_str: Formatted date string for filename (e.g., "2.3")

    Returns:
        Tuple of (success, message, new_filename or None)
    """
    try:
        # Load the document
        doc = Document(doc_path)

        # Extract all non-empty paragraphs
        paragraphs = [p for p in doc.paragraphs if p.text.strip()]

        if len(paragraphs) < 3:
            return False, "Document has less than 3 lines of text", None

        # Get the 3rd line (index 2) before deletion
        third_line = paragraphs[2].text.strip()

        # Delete the first line (remove the first non-empty paragraph entirely)
        for para in doc.paragraphs:
            if para.text.strip():
                # Remove the paragraph element from the document
                p = para._element
                p.getparent().remove(p)
                break

        # Save the modified document
        doc.save(doc_path)

        # Create new filename
        subfolder_part = sanitize_for_filename(subfolder_name)
        third_line_part = sanitize_for_filename(third_line)

        new_filename = f"{target_date_str}_{subfolder_part}_{third_line_part}.docx"

        return True, "Processed successfully", new_filename

    except Exception as e:
        return False, f"Error: {str(e)}", None


def clean_and_rename_files(
    root_directory: Path,
    progress_callback=None
) -> Tuple[int, int]:
    """
    Process all Word documents in subfolders of the root directory.

    Args:
        root_directory: Directory containing subfolders with Word files
        progress_callback: Optional callback function(message: str)

    Returns:
        Tuple of (files_processed, files_skipped)
    """
    def log(msg):
        if progress_callback:
            progress_callback(msg)
        else:
            print(msg)

    if not root_directory.exists():
        log(f"Error: Directory '{root_directory}' does not exist.")
        return 0, 0

    # Calculate the target date (2 Mondays from now)
    target_date = get_two_mondays_from_now()
    target_date_str = format_date_for_filename(target_date)

    log(f"Target date: {target_date.strftime('%m/%d/%Y')} ({target_date_str})")

    processed_count = 0
    skipped_count = 0

    # Iterate through all subdirectories
    for subfolder in sorted(root_directory.iterdir()):
        if not subfolder.is_dir():
            continue

        subfolder_name = subfolder.name
        log(f"Processing subfolder: {subfolder_name}")

        # Find the Word file matching the subfolder name
        word_file = find_matching_word_file(subfolder, subfolder_name)

        if not word_file:
            log(f"  No matching Word file found. Skipping.")
            skipped_count += 1
            continue

        log(f"  Found file: {word_file.name}")

        # Process the document
        success, message, new_filename = process_single_document(
            word_file, subfolder_name, target_date_str
        )

        if success and new_filename:
            # Rename the file
            new_path = subfolder / new_filename

            if new_path.exists() and new_path != word_file:
                log(f"  Warning: Target filename already exists: {new_filename}")
                log(f"  Overwriting existing file.")
                new_path.unlink()

            word_file.rename(new_path)
            log(f"  Renamed to: {new_filename}")
            processed_count += 1
        else:
            log(f"  {message}. Skipping.")
            skipped_count += 1

    return processed_count, skipped_count
