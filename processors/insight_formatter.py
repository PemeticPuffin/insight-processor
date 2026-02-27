"""
Step 3: Format Insight Files

Applies formatting rules to insight documents:
- Set font to Calibri 11pt throughout
- Bold specific headers
- Convert all content in bullet sections to dot bullets
- Ensure spacing after "Title of Insight"
- Convert AskGartner prompts to hyperlinks
- Replace CXO placeholder with the role's full name
- Preserve all hyperlinks
"""

from pathlib import Path
from docx import Document
from typing import Tuple, List, Set

import sys
sys.path.insert(0, str(Path(__file__).parent.parent))

from config import HEADERS_TO_BOLD, BULLET_SECTION_HEADERS, ASKGARTNER_HEADER, ASKGARTNER_BASE_URL
from utils.formatting_utils import (
    set_font_calibri_11pt,
    bold_paragraph,
    apply_bullet_style,
    apply_bullet_with_hyperlink,
    add_spacing_after,
    replace_cxo_in_document,
)
from utils.file_utils import find_word_files
from utils.text_utils import get_role_full_name


def is_header(paragraph) -> bool:
    """Check if paragraph text matches one of our headers."""
    text = paragraph.text.strip()
    return text in HEADERS_TO_BOLD


def get_section_ranges(doc: Document) -> List[Tuple[int, int, int]]:
    """
    Identify paragraph indices for sections that need bullet points.

    For each header in BULLET_SECTION_HEADERS, bullets are applied to content
    from that header until the next header in HEADERS_TO_BOLD (or end of document).

    Args:
        doc: The Document to analyze

    Returns:
        List of (start_idx, end_idx, first_content_idx) tuples for bullet sections
        first_content_idx is the index of the first non-empty paragraph in the section
    """
    ranges = []

    # Find positions of all recognized headers
    header_positions = []  # List of (idx, header_text) tuples
    for idx, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        if text in HEADERS_TO_BOLD:
            header_positions.append((idx, text))

    # Sort by position (should already be sorted, but be safe)
    header_positions.sort(key=lambda x: x[0])

    # For each header that should have bulleted content
    for i, (idx, header_text) in enumerate(header_positions):
        if header_text in BULLET_SECTION_HEADERS:
            start_idx = idx + 1  # Start after the header

            # Find the next header position (or end of document)
            if i + 1 < len(header_positions):
                end_idx = header_positions[i + 1][0]
            else:
                end_idx = len(doc.paragraphs)

            # Find the first non-empty paragraph in this section
            first_content_idx = -1
            for j in range(start_idx, end_idx):
                if doc.paragraphs[j].text.strip():
                    first_content_idx = j
                    break

            if start_idx < end_idx:
                ranges.append((start_idx, end_idx, first_content_idx))

    return ranges


def get_askgartner_indices(doc: Document) -> Set[int]:
    """
    Identify paragraph indices that belong to the AskGartner section.

    Args:
        doc: The Document to analyze

    Returns:
        Set of paragraph indices within the AskGartner section
    """
    indices = set()
    in_askgartner = False

    for idx, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        if text == ASKGARTNER_HEADER:
            in_askgartner = True
            continue
        elif text in HEADERS_TO_BOLD:
            in_askgartner = False
            continue

        if in_askgartner:
            indices.add(idx)

    return indices


def find_title_paragraph(doc: Document) -> int:
    """
    Find the paragraph index of "Title of Insight" or similar title line.

    Args:
        doc: The Document to search

    Returns:
        Paragraph index, or -1 if not found
    """
    for idx, para in enumerate(doc.paragraphs):
        text = para.text.strip().lower()
        if 'title of insight' in text or text.startswith('title:'):
            return idx
    return -1


def format_single_document(file_path: Path, full_name: str = None) -> Tuple[bool, str]:
    """
    Apply all formatting rules to a single document.

    Args:
        file_path: Path to the document to format
        full_name: Optional role full name; if provided, replaces 'CXO' with this value

    Returns:
        Tuple of (success, message)
    """
    try:
        doc = Document(file_path)

        # Get section ranges that need bullet points
        bullet_ranges = get_section_ranges(doc)

        # Create a set of paragraph indices that should be bulleted
        bullet_indices: Set[int] = set()
        for start, end, _ in bullet_ranges:
            bullet_indices.update(range(start, end))

        # Identify AskGartner section paragraphs for hyperlink conversion
        askgartner_indices = get_askgartner_indices(doc)

        # Find title paragraph for spacing
        title_idx = find_title_paragraph(doc)

        # Process each paragraph
        for idx, para in enumerate(doc.paragraphs):
            # Set font to Calibri 11pt
            set_font_calibri_11pt(para)

            # Bold headers
            if is_header(para):
                bold_paragraph(para)

            # Add spacing after title
            elif idx == title_idx:
                add_spacing_after(para)

            # Handle bullet sections
            elif idx in bullet_indices and para.text.strip():
                if idx in askgartner_indices:
                    apply_bullet_with_hyperlink(para, ASKGARTNER_BASE_URL, doc)
                else:
                    apply_bullet_style(para)

        # Replace CXO placeholder with role's full name
        if full_name:
            replace_cxo_in_document(doc, full_name)

        # Save the document
        doc.save(file_path)
        return True, "Formatted successfully"

    except Exception as e:
        return False, f"Error: {str(e)}"


def format_insights(
    directory: Path,
    recursive: bool = True,
    progress_callback=None
) -> Tuple[int, int]:
    """
    Format all Word documents in a directory.

    Args:
        directory: Directory containing files to format
        recursive: If True, search subdirectories
        progress_callback: Optional callback function(message: str)

    Returns:
        Tuple of (success_count, failure_count)
    """
    def log(msg):
        if progress_callback:
            progress_callback(msg)
        else:
            print(msg)

    # Find all Word documents
    word_files = find_word_files(directory, recursive=recursive)

    if not word_files:
        log("No Word documents found.")
        return 0, 0

    log(f"Found {len(word_files)} Word document(s)")

    success_count = 0
    failure_count = 0

    for file_path in word_files:
        # Skip email files generated by Step 4
        if file_path.name.endswith('_BD_AE_emails.docx') or file_path.name.endswith('_EP_email.docx'):
            log(f"  Skipping email file: {file_path.name}")
            continue

        log(f"Processing: {file_path.name}")

        # Resolve full name from the file's parent subfolder for CXO replacement
        full_name = get_role_full_name(file_path.parent.name)

        success, message = format_single_document(file_path, full_name=full_name)

        if success:
            log(f"  {message}")
            success_count += 1
        else:
            log(f"  {message}")
            failure_count += 1

    return success_count, failure_count
